"""DiligenceAI — Deal Analysis Platform. Single-file Streamlit app."""
import streamlit as st
import io, os, csv, json, re
import urllib.request, urllib.error
import sqlite3, hashlib, uuid
from datetime import datetime

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    EXCEL_OK = True
except ImportError:
    EXCEL_OK = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    PDF_OK = True
except ImportError:
    PDF_OK = False

# ─────────────────────────────────────────────────────────────────────────────
# DATABASE  (SQLite, fully inlined)
# ─────────────────────────────────────────────────────────────────────────────
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "diligenceai.db")

def _db():
    c = sqlite3.connect(DB_PATH, check_same_thread=False)
    c.row_factory = sqlite3.Row
    return c

def _init_db():
    db = _db()
    db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id            TEXT PRIMARY KEY,
            email         TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at    TEXT DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS deals (
            id              TEXT PRIMARY KEY,
            user_id         TEXT NOT NULL,
            name            TEXT DEFAULT 'Unnamed Deal',
            industry        TEXT DEFAULT '',
            status          TEXT DEFAULT 'screening',
            screening_data  TEXT DEFAULT '{}',
            dd_data         TEXT DEFAULT '{}',
            risk_notes      TEXT DEFAULT '{}',
            file_names      TEXT DEFAULT '[]',
            created_at      TEXT DEFAULT (datetime('now')),
            updated_at      TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS analyses (
            id            TEXT PRIMARY KEY,
            user_id       TEXT NOT NULL,
            deal_id       TEXT DEFAULT '',
            company_name  TEXT DEFAULT 'Unknown Company',
            period        TEXT DEFAULT '',
            health_score  INTEGER DEFAULT 5,
            health_label  TEXT DEFAULT 'Moderate',
            key_metrics   TEXT DEFAULT '{}',
            insights      TEXT DEFAULT '{}',
            raw_output    TEXT DEFAULT '{}',
            date_created  TEXT DEFAULT (datetime('now')),
            FOREIGN KEY (user_id) REFERENCES users(id)
        );
        CREATE TABLE IF NOT EXISTS shared_reports (
            share_id      TEXT PRIMARY KEY,
            user_id       TEXT NOT NULL,
            company_name  TEXT,
            period        TEXT,
            health_score  INTEGER,
            health_label  TEXT,
            raw_output    TEXT,
            created_at    TEXT DEFAULT (datetime('now'))
        );
    """)
    db.commit()
    db.close()

def _hash(pw: str) -> str:
    return hashlib.sha256(pw.encode("utf-8")).hexdigest()

def db_create_user(email: str, password: str) -> dict:
    db = _db()
    try:
        uid = str(uuid.uuid4())
        db.execute(
            "INSERT INTO users (id, email, password_hash) VALUES (?,?,?)",
            (uid, email.lower().strip(), _hash(password))
        )
        db.commit()
        return {"ok": True, "user_id": uid, "email": email.lower().strip()}
    except sqlite3.IntegrityError:
        return {"ok": False, "error": "An account with this email already exists."}
    finally:
        db.close()

def db_login_user(email: str, password: str) -> dict:
    db = _db()
    row = db.execute(
        "SELECT * FROM users WHERE email=? AND password_hash=?",
        (email.lower().strip(), _hash(password))
    ).fetchone()
    db.close()
    if row:
        return {"ok": True, "user_id": row["id"], "email": row["email"]}
    return {"ok": False, "error": "Incorrect email or password."}

# ── DEAL CRUD ──────────────────────────────────────────────────────────────
def db_create_deal(user_id: str, name: str, industry: str = "", file_names: list = None) -> str:
    db = _db()
    did = str(uuid.uuid4())
    db.execute(
        "INSERT INTO deals (id, user_id, name, industry, file_names) VALUES (?,?,?,?,?)",
        (did, user_id, name, industry, json.dumps(file_names or []))
    )
    db.commit()
    db.close()
    return did

def db_save_screening(deal_id: str, screening_data: dict, file_names: list = None):
    db = _db()
    db.execute(
        "UPDATE deals SET screening_data=?, status='screening', updated_at=datetime('now'), file_names=? WHERE id=?",
        (json.dumps(screening_data), json.dumps(file_names or []), deal_id)
    )
    db.commit()
    db.close()

def db_promote_to_dd(deal_id: str):
    db = _db()
    db.execute(
        "UPDATE deals SET status='dd', updated_at=datetime('now') WHERE id=?",
        (deal_id,)
    )
    db.commit()
    db.close()

def db_save_dd(deal_id: str, dd_data: dict):
    db = _db()
    db.execute(
        "UPDATE deals SET dd_data=?, updated_at=datetime('now') WHERE id=?",
        (json.dumps(dd_data), deal_id)
    )
    db.commit()
    db.close()

def db_save_risk_notes(deal_id: str, notes: dict):
    db = _db()
    db.execute(
        "UPDATE deals SET risk_notes=?, updated_at=datetime('now') WHERE id=?",
        (json.dumps(notes), deal_id)
    )
    db.commit()
    db.close()

def db_get_deals(user_id: str) -> list:
    db = _db()
    rows = db.execute(
        "SELECT * FROM deals WHERE user_id=? ORDER BY updated_at DESC", (user_id,)
    ).fetchall()
    db.close()
    result = []
    for r in rows:
        d = dict(r)
        d["screening_data"] = json.loads(d.get("screening_data") or "{}")
        d["dd_data"]        = json.loads(d.get("dd_data") or "{}")
        d["risk_notes"]     = json.loads(d.get("risk_notes") or "{}")
        d["file_names"]     = json.loads(d.get("file_names") or "[]")
        result.append(d)
    return result

def db_get_deal(deal_id: str):
    db = _db()
    row = db.execute("SELECT * FROM deals WHERE id=?", (deal_id,)).fetchone()
    db.close()
    if not row:
        return None
    d = dict(row)
    d["screening_data"] = json.loads(d.get("screening_data") or "{}")
    d["dd_data"]        = json.loads(d.get("dd_data") or "{}")
    d["risk_notes"]     = json.loads(d.get("risk_notes") or "{}")
    d["file_names"]     = json.loads(d.get("file_names") or "[]")
    return d

def db_delete_deal(deal_id: str, user_id: str) -> bool:
    db = _db()
    cur = db.execute("DELETE FROM deals WHERE id=? AND user_id=?", (deal_id, user_id))
    db.commit()
    db.close()
    return cur.rowcount > 0

# ── ANALYSES (DD reports) ──────────────────────────────────────────────────
def db_save_analysis(user_id: str, data: dict, deal_id: str = "") -> str:
    db = _db()
    aid = str(uuid.uuid4())
    kpis = data.get("kpis", {})
    key_metrics = json.dumps({
        k: kpis.get(k, {}).get("value", "N/A")
        for k in ["revenue","net_profit","gross_margin","net_margin","ebitda",
                  "operating_cashflow","current_ratio","debt_to_equity",
                  "working_capital","total_debt","revenue_growth","interest_coverage"]
    })
    insights = json.dumps({
        "health_summary":  data.get("health_summary", ""),
        "investor_view":   data.get("investor_view", ""),
        "risks":           data.get("risks", []),
        "positives":       data.get("positives", []),
        "recommendations": data.get("recommendations", []),
    })
    db.execute(
        """INSERT INTO analyses
           (id, user_id, deal_id, company_name, period, health_score, health_label,
            key_metrics, insights, raw_output)
           VALUES (?,?,?,?,?,?,?,?,?,?)""",
        (aid, user_id, deal_id,
         data.get("company_name","Unknown Company"),
         data.get("period",""),
         data.get("health_score",5),
         data.get("health_label","Moderate"),
         key_metrics, insights, json.dumps(data))
    )
    db.commit()
    db.close()
    return aid

def db_get_analyses(user_id: str) -> list:
    db = _db()
    rows = db.execute(
        "SELECT * FROM analyses WHERE user_id=? ORDER BY date_created DESC", (user_id,)
    ).fetchall()
    db.close()
    result = []
    for r in rows:
        d = dict(r)
        d["key_metrics"] = json.loads(d.get("key_metrics") or "{}")
        d["raw_output"]  = json.loads(d.get("raw_output") or "{}")
        result.append(d)
    return result

def db_get_analysis(analysis_id: str):
    db = _db()
    row = db.execute("SELECT * FROM analyses WHERE id=?", (analysis_id,)).fetchone()
    db.close()
    if not row:
        return None
    d = dict(row)
    d["key_metrics"] = json.loads(d.get("key_metrics") or "{}")
    d["raw_output"]  = json.loads(d.get("raw_output") or "{}")
    return d

def db_delete_analysis(analysis_id: str, user_id: str) -> bool:
    db = _db()
    cur = db.execute("DELETE FROM analyses WHERE id=? AND user_id=?", (analysis_id, user_id))
    db.commit()
    db.close()
    return cur.rowcount > 0

def db_create_share(user_id: str, data: dict) -> str:
    db = _db()
    share_id = str(uuid.uuid4())[:8].upper()
    db.execute(
        """INSERT INTO shared_reports
           (share_id, user_id, company_name, period, health_score, health_label, raw_output)
           VALUES (?,?,?,?,?,?,?)""",
        (share_id, user_id,
         data.get("company_name","Unknown Company"),
         data.get("period",""),
         data.get("health_score",5),
         data.get("health_label","Moderate"),
         json.dumps(data))
    )
    db.commit()
    db.close()
    return share_id

def db_get_shared(share_id: str):
    db = _db()
    row = db.execute(
        "SELECT * FROM shared_reports WHERE share_id=?", (share_id.upper().strip(),)
    ).fetchone()
    db.close()
    if not row:
        return None
    d = dict(row)
    d["raw_output"] = json.loads(d.get("raw_output") or "{}")
    return d

_init_db()

# ─────────────────────────────────────────────────────────────────────────────
# PDF BUILDER
# ─────────────────────────────────────────────────────────────────────────────
def build_pdf(data: dict) -> bytes:
    if not PDF_OK:
        return b""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    C_BG   = rl_colors.HexColor("#080C14")
    C_CARD = rl_colors.HexColor("#0F1620")
    C_BLUE = rl_colors.HexColor("#4F8EF7")
    C_TEAL = rl_colors.HexColor("#00D4AA")
    C_PURP = rl_colors.HexColor("#9B6DFF")
    C_AMB  = rl_colors.HexColor("#F5A623")
    C_RED  = rl_colors.HexColor("#FF5C6A")
    C_SUB  = rl_colors.HexColor("#8B9BC8")
    C_MUT  = rl_colors.HexColor("#4A5578")
    C_BOR  = rl_colors.HexColor("#1A2340")
    C_TXT  = rl_colors.HexColor("#F0F4FF")
    def S(name, **kw):
        base = dict(fontName="Helvetica", fontSize=10, textColor=C_TXT, leading=14, spaceAfter=4)
        base.update(kw)
        return ParagraphStyle(name, **base)
    lbl   = data.get("health_label","Moderate")
    score = data.get("health_score", 5)
    hcol  = {"Strong": C_TEAL, "Moderate": C_AMB, "Weak": C_RED}.get(lbl, C_AMB)
    div   = HRFlowable(width="100%", thickness=1, color=C_BOR, spaceAfter=10)
    story = []
    story.append(Paragraph("FINANCIAL STATEMENT ANALYSIS",
                            S("ti", fontSize=22, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)))
    story.append(Paragraph("DiligenceAI — Deal Analysis Platform",
                            S("su", fontSize=11, textColor=C_SUB, alignment=TA_CENTER, spaceAfter=16)))
    story.append(div)
    story.append(Paragraph(f"{data.get('company_name','Unknown')}  ·  {data.get('period','')}",
                            S("co", fontSize=14, fontName="Helvetica-Bold", spaceAfter=2)))
    story.append(Paragraph(f"Financial Health: {lbl}  |  Score: {score} / 10",
                            S("hs", fontSize=12, fontName="Helvetica-Bold", textColor=hcol, spaceAfter=12)))
    story.append(div)
    story.append(Paragraph("EXECUTIVE SUMMARY", S("lb", fontSize=7, fontName="Helvetica-Bold", textColor=C_MUT, spaceAfter=4)))
    story.append(Paragraph(data.get("health_summary",""), S("bo", fontSize=10, textColor=C_SUB, leading=15, spaceAfter=8)))
    story.append(Paragraph("INVESTOR VIEW", S("lb2", fontSize=7, fontName="Helvetica-Bold", textColor=C_MUT, spaceAfter=4)))
    story.append(Paragraph(data.get("investor_view",""), S("iv", fontSize=10, textColor=C_SUB, leading=15, spaceAfter=8)))
    story.append(div)
    story.append(Paragraph("KEY FINANCIAL METRICS",
                            S("h2", fontSize=13, fontName="Helvetica-Bold", textColor=C_BLUE, spaceBefore=6, spaceAfter=6)))
    kpi_keys = [("revenue","Revenue"),("net_profit","Net Profit"),("gross_margin","Gross Margin"),
                ("net_margin","Net Margin"),("ebitda","EBITDA"),("operating_cashflow","Operating Cash Flow"),
                ("current_ratio","Current Ratio"),("debt_to_equity","Debt / Equity"),
                ("working_capital","Working Capital"),("total_debt","Total Debt"),
                ("revenue_growth","Revenue Growth"),("interest_coverage","Interest Coverage")]
    kpis  = data.get("kpis", {})
    tdata = [["Metric","Value","Commentary"]]
    for key, label in kpi_keys:
        item = kpis.get(key, {})
        tdata.append([label, item.get("value","N/A"), item.get("note","")])
    t = Table(tdata, colWidths=[5*cm, 4*cm, 8*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_BLUE), ("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"), ("FONTSIZE",(0,0),(-1,-1),9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_CARD,C_BG]),
        ("TEXTCOLOR",(0,1),(-1,-1),C_SUB), ("TEXTCOLOR",(1,1),(1,-1),C_BLUE),
        ("FONTNAME",(1,1),(1,-1),"Helvetica-Bold"),
        ("GRID",(0,0),(-1,-1),0.5,C_BOR),
        ("TOPPADDING",(0,0),(-1,-1),5), ("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),8),
    ]))
    story.append(t); story.append(Spacer(1,12)); story.append(div)
    for sec_key, sec_title, accent in [
        ("profitability","Profitability",C_BLUE), ("cash_health","Cash Health",C_TEAL),
        ("revenue_growth","Revenue Growth",C_PURP), ("working_capital_analysis","Working Capital",C_AMB),
        ("balance_sheet","Balance Sheet",C_RED), ("debt_leverage","Debt & Leverage",C_TEAL)]:
        sec = data.get(sec_key, {})
        story.append(Paragraph(sec_title.upper(),
                                S(f"sh{sec_key}", fontSize=13, fontName="Helvetica-Bold", textColor=accent, spaceBefore=10, spaceAfter=4)))
        story.append(Paragraph(sec.get("headline",""),
                                S(f"hd{sec_key}", fontSize=10, textColor=C_SUB, leading=15, spaceAfter=4)))
        for pt in sec.get("points",[]):
            story.append(Paragraph(f"• {pt}", S(f"pt{sec_key}", fontSize=10, textColor=C_SUB, leading=14, spaceAfter=3)))
        story.append(Spacer(1,4))
    story.append(div)
    story.append(Paragraph("KEY RISKS & CONCERNS",
                            S("rh", fontSize=13, fontName="Helvetica-Bold", textColor=C_RED, spaceBefore=6, spaceAfter=6)))
    for i, risk in enumerate(data.get("risks",[])):
        story.append(Paragraph(f"<b>{risk.get('title','')}</b>", S(f"rt{i}", fontSize=10, textColor=C_RED, spaceAfter=2)))
        story.append(Paragraph(f"Issue: {risk.get('detail','')}", S(f"rd{i}", fontSize=10, textColor=C_SUB, leading=14, spaceAfter=2)))
        story.append(Paragraph(f"Action: {risk.get('fix','')}", S(f"rf{i}", fontSize=10, textColor=C_TEAL, leading=14, spaceAfter=6)))
    story.append(Paragraph("POSITIVE SIGNALS",
                            S("ph", fontSize=13, fontName="Helvetica-Bold", textColor=C_TEAL, spaceBefore=6, spaceAfter=6)))
    for i, pos in enumerate(data.get("positives",[])):
        story.append(Paragraph(f"<b>{pos.get('title','')}</b>", S(f"pos{i}", fontSize=10, textColor=C_TEAL, spaceAfter=2)))
        story.append(Paragraph(pos.get("detail",""), S(f"posd{i}", fontSize=10, textColor=C_SUB, leading=14, spaceAfter=6)))
    story.append(Paragraph("RECOMMENDATIONS",
                            S("rch", fontSize=13, fontName="Helvetica-Bold", textColor=C_BLUE, spaceBefore=6, spaceAfter=6)))
    for i, rec in enumerate(data.get("recommendations",[]), 1):
        story.append(Paragraph(f"<b>{i}. {rec.get('action','')}</b>", S(f"rc{i}", fontSize=10, textColor=C_BLUE, spaceAfter=2)))
        story.append(Paragraph(rec.get("rationale",""), S(f"rcd{i}", fontSize=10, textColor=C_SUB, leading=14, spaceAfter=6)))
    story.append(div)
    story.append(Paragraph("DiligenceAI  ·  For informational purposes only — not financial advice.",
                            S("ft", fontSize=8, textColor=C_MUT, alignment=TA_CENTER)))
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
# EXCEL BUILDER
# ─────────────────────────────────────────────────────────────────────────────
def build_excel(data: dict) -> bytes:
    if not EXCEL_OK:
        return b""
    wb = Workbook()
    DN="080C14"; AB="4F8EF7"; LB="0F2040"; WH="FFFFFF"; MG="8B9BC8"
    GB="0A2018"; GF="00D4AA"; RB="200A0C"; RF="FF5C6A"; AMB="201408"; AMF="F5A623"; MN="0F1620"
    lbl=data.get("health_label","Moderate"); sc=data.get("health_score",5)
    hc={"Strong":GF,"Moderate":AMF,"Weak":RF}.get(lbl,AMF)
    hb={"Strong":GB,"Moderate":AMB,"Weak":RB}.get(lbl,AMB)
    def hf(sz=11,b=True,c=WH): return Font(name="Arial",size=sz,bold=b,color=c)
    def bf(sz=10,b=False,c="D0D8F0"): return Font(name="Arial",size=sz,bold=b,color=c)
    def fl(h): return PatternFill("solid",fgColor=h)
    def tb(s="all"):
        sd=Side(style="thin",color="1A2340"); n=Side(style=None)
        return Border(left=sd if "all" in s or "left" in s else n,
                      right=sd if "all" in s or "right" in s else n,
                      top=sd if "all" in s or "top" in s else n,
                      bottom=sd if "all" in s or "bottom" in s else n)
    def ca(): return Alignment(horizontal="center",vertical="center",wrap_text=True)
    def la(w=True): return Alignment(horizontal="left",vertical="center",wrap_text=w)
    def rh(ws,r,h): ws.row_dimensions[r].height=h
    def mw(ws,rng,v,fn,al,fl_=None):
        ws.merge_cells(rng); c=ws[rng.split(":")[0]]
        c.value=v; c.font=fn; c.alignment=al
        if fl_: c.fill=fl_
    ws=wb.active; ws.title="Executive Summary"; ws.sheet_view.showGridLines=False
    for i,w in enumerate([2,28,22,22,22,22,2],1): ws.column_dimensions[get_column_letter(i)].width=w
    r=1
    for ri in range(r,r+3):
        rh(ws,ri,6 if ri!=r+1 else 40)
        for ci in range(1,8): ws.cell(ri,ci).fill=fl(DN)
    mw(ws,f"B{r+1}:F{r+1}","FINANCIAL STATEMENT ANALYSIS — DiligenceAI",hf(15),ca(),fl(DN)); r+=3
    for ri in range(r,r+4):
        rh(ws,ri,6 if ri in(r,r+3) else 28)
        for ci in range(1,8): ws.cell(ri,ci).fill=fl(MN)
    ws.merge_cells(f"B{r+1}:C{r+1}")
    c=ws[f"B{r+1}"]; c.value=data.get("company_name","Unknown"); c.font=hf(13); c.alignment=la(False)
    ws.cell(r+1,4).value=data.get("period",""); ws.cell(r+1,4).font=hf(11,False,MG); ws.cell(r+1,4).alignment=ca()
    for ci,val in enumerate([lbl,f"Score:{sc}/10"],5):
        c=ws.cell(r+1,ci); c.value=val; c.font=Font(name="Arial",size=11,bold=True,color=hc)
        c.fill=fl(hb); c.alignment=ca(); c.border=tb()
    r+=4; rh(ws,r,6); r+=1; rh(ws,r,18)
    mw(ws,f"B{r}:F{r}","EXECUTIVE SUMMARY",hf(10),la(False),fl(AB)); r+=1
    for line in data.get("health_summary","").split(". "):
        if not line.strip(): continue
        rh(ws,r,32); ws.merge_cells(f"B{r}:F{r}"); c=ws[f"B{r}"]
        c.value=line.strip().rstrip(".")+"."; c.font=bf(10); c.alignment=la(); c.fill=fl(MN); c.border=tb("bottom"); r+=1
    rh(ws,r,6); r+=1; rh(ws,r,18)
    mw(ws,f"B{r}:F{r}","INVESTOR VIEW",hf(10),la(False),fl(AB)); r+=1
    iv=data.get("investor_view",""); rh(ws,r,max(60,len(iv)//4))
    ws.merge_cells(f"B{r}:F{r}"); c=ws[f"B{r}"]
    c.value=iv; c.font=bf(10,b=True,c="C8D0E8"); c.alignment=la(); c.fill=fl(LB); c.border=tb()
    ws2=wb.create_sheet("KPI Metrics"); ws2.sheet_view.showGridLines=False
    for i,w in enumerate([2,30,20,35,2],1): ws2.column_dimensions[get_column_letter(i)].width=w
    r=1
    for ri in range(r,r+3):
        rh(ws2,ri,6 if ri!=r+1 else 36)
        for ci in range(1,6): ws2.cell(ri,ci).fill=fl(DN)
    mw(ws2,f"B{r+1}:D{r+1}","KEY FINANCIAL METRICS",hf(14),ca(),fl(DN)); r+=3
    rh(ws2,r,20)
    for ci,h in enumerate(["Metric","Value","Commentary"],2):
        c=ws2.cell(r,ci); c.value=h; c.font=hf(10); c.fill=fl(AB); c.alignment=ca(); c.border=tb()
    r+=1; kpis=data.get("kpis",{})
    for idx,(key,label) in enumerate([("revenue","Revenue"),("net_profit","Net Profit"),
        ("gross_margin","Gross Margin"),("net_margin","Net Margin"),("ebitda","EBITDA"),
        ("operating_cashflow","Operating Cash Flow"),("current_ratio","Current Ratio"),
        ("debt_to_equity","Debt/Equity"),("working_capital","Working Capital"),("total_debt","Total Debt"),
        ("revenue_growth","Revenue Growth"),("interest_coverage","Interest Coverage")]):
        item=kpis.get(key,{}); rf=fl(MN) if idx%2==0 else fl(DN); rh(ws2,r,22)
        c=ws2.cell(r,2); c.value=label; c.font=bf(10,b=True,c="F0F4FF"); c.fill=rf; c.alignment=la(False); c.border=tb()
        c=ws2.cell(r,3); c.value=item.get("value","N/A"); c.font=Font(name="Arial",size=11,bold=True,color=AB); c.fill=rf; c.alignment=ca(); c.border=tb()
        c=ws2.cell(r,4); c.value=item.get("note",""); c.font=bf(9,c="8B9BC8"); c.fill=rf; c.alignment=la(); c.border=tb(); r+=1
    buf=io.BytesIO(); wb.save(buf); buf.seek(0); return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
# WORD BUILDER
# ─────────────────────────────────────────────────────────────────────────────
def build_docx(data: dict) -> bytes:
    if not DOCX_OK:
        return b""
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Inches(1); sec.bottom_margin=Inches(1)
        sec.left_margin=Inches(1.2); sec.right_margin=Inches(1.2)
    def ah(text, lvl=1, col=None):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        run=p.add_run(text); run.bold=True; run.font.name="Arial"
        run.font.size=Pt(16 if lvl==1 else 13 if lvl==2 else 11)
        if col: run.font.color.rgb=RGBColor(*col)
        return p
    def ab(text, italic=False, col=None):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
        run=p.add_run(text); run.font.name="Arial"; run.font.size=Pt(10); run.italic=italic
        if col: run.font.color.rgb=RGBColor(*col)
        return p
    def abul(text):
        p=doc.add_paragraph(style="List Bullet")
        run=p.add_run(text); run.font.name="Arial"; run.font.size=Pt(10)
    def adiv():
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
        run=p.add_run("─"*60); run.font.color.rgb=RGBColor(42,52,80); run.font.size=Pt(8)
    lbl=data.get("health_label","Moderate"); score=data.get("health_score",5)
    hcol={"Strong":(0,212,170),"Moderate":(245,166,35),"Weak":(255,92,106)}.get(lbl,(245,166,35))
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr=t.add_run("FINANCIAL STATEMENT ANALYSIS"); tr.bold=True; tr.font.name="Arial"; tr.font.size=Pt(20)
    s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sr=s.add_run("DiligenceAI — Deal Analysis Platform")
    sr.font.name="Arial"; sr.font.size=Pt(11); sr.font.color.rgb=RGBColor(139,155,200)
    adiv(); doc.add_paragraph()
    cp=doc.add_paragraph(); cr=cp.add_run(f"{data.get('company_name','Unknown')}  ·  {data.get('period','')}")
    cr.bold=True; cr.font.name="Arial"; cr.font.size=Pt(13)
    hp=doc.add_paragraph(); hr=hp.add_run(f"Financial Health: {lbl}   |   Score: {score} / 10")
    hr.bold=True; hr.font.name="Arial"; hr.font.size=Pt(12); hr.font.color.rgb=RGBColor(*hcol)
    adiv(); ah("Executive Summary",2,(79,142,247)); ab(data.get("health_summary",""))
    ah("Investor View",2,(79,142,247)); ab(data.get("investor_view",""),italic=True); adiv()
    ah("Key Financial Metrics",2,(79,142,247))
    kpi_labels=[("revenue","Revenue"),("net_profit","Net Profit"),("gross_margin","Gross Margin"),
                ("net_margin","Net Margin"),("ebitda","EBITDA"),("operating_cashflow","Operating Cash Flow"),
                ("current_ratio","Current Ratio"),("debt_to_equity","Debt / Equity"),
                ("working_capital","Working Capital"),("total_debt","Total Debt")]
    kpis=data.get("kpis",{}); tbl=doc.add_table(rows=1,cols=3); tbl.style="Table Grid"
    hdr=tbl.rows[0].cells
    for i,h in enumerate(["Metric","Value","Commentary"]):
        hdr[i].text=h; run=hdr[i].paragraphs[0].runs[0]; run.bold=True; run.font.name="Arial"; run.font.size=Pt(10)
    for key,label in kpi_labels:
        item=kpis.get(key,{}); row=tbl.add_row().cells
        row[0].text=label; row[1].text=item.get("value","N/A"); row[2].text=item.get("note","")
        for cell in row:
            for para in cell.paragraphs:
                for r in para.runs: r.font.name="Arial"; r.font.size=Pt(10)
    adiv()
    for sk,sl in [("profitability","Profitability"),("cash_health","Cash Health"),
                  ("revenue_growth","Revenue Growth"),("working_capital_analysis","Working Capital"),
                  ("balance_sheet","Balance Sheet"),("debt_leverage","Debt & Leverage")]:
        sec=data.get(sk,{}); ah(sl,2,(79,142,247)); ab(sec.get("headline",""),italic=True,col=(139,155,200))
        for pt in sec.get("points",[]): abul(pt)
    adiv()
    ah("Key Risks & Concerns",2,(255,92,106))
    for risk in data.get("risks",[]):
        rp=doc.add_paragraph(); rr=rp.add_run(risk.get("title",""))
        rr.bold=True; rr.font.name="Arial"; rr.font.size=Pt(10); rr.font.color.rgb=RGBColor(255,92,106)
        ab(f"Issue: {risk.get('detail','')}"); ab(f"Action: {risk.get('fix','')}",col=(0,212,170))
    ah("Positive Signals",2,(0,212,170))
    for pos in data.get("positives",[]):
        pp=doc.add_paragraph(); pr=pp.add_run(pos.get("title",""))
        pr.bold=True; pr.font.name="Arial"; pr.font.size=Pt(10); pr.font.color.rgb=RGBColor(0,212,170)
        ab(pos.get("detail",""))
    ah("Recommendations",2,(79,142,247))
    for i,rec in enumerate(data.get("recommendations",[]),1):
        rp=doc.add_paragraph(); rr=rp.add_run(f"{i}. {rec.get('action','')}")
        rr.bold=True; rr.font.name="Arial"; rr.font.size=Pt(10); rr.font.color.rgb=RGBColor(79,142,247)
        ab(rec.get("rationale",""),col=(139,155,200))
    adiv()
    fn=doc.add_paragraph(); fn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fr=fn.add_run("DiligenceAI  ·  For informational purposes only — not financial advice.")
    fr.font.name="Arial"; fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(74,85,120)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()

# ─────────────────────────────────────────────────────────────────────────────
# TXT BUILDER
# ─────────────────────────────────────────────────────────────────────────────
def build_txt(data: dict) -> str:
    kpi_map = {"revenue":"Revenue","net_profit":"Net Profit","gross_margin":"Gross Margin",
               "net_margin":"Net Margin","ebitda":"EBITDA","operating_cashflow":"Operating Cash Flow",
               "current_ratio":"Current Ratio","debt_to_equity":"Debt/Equity",
               "working_capital":"Working Capital","total_debt":"Total Debt"}
    kpis = data.get("kpis",{})
    lines = ["FINANCIAL ANALYSIS REPORT — DiligenceAI","="*60,
             f"Company : {data.get('company_name','N/A')}",
             f"Period  : {data.get('period','N/A')}",
             f"Health  : {data.get('health_label','N/A')}   Score: {data.get('health_score','N/A')}/10",
             "","EXECUTIVE SUMMARY","-"*40,data.get("health_summary",""),
             "","INVESTOR VIEW","-"*40,data.get("investor_view",""),
             "","KEY METRICS","-"*40]
    for k,lb in kpi_map.items():
        v=kpis.get(k,{}).get("value","N/A"); n=kpis.get(k,{}).get("note","")
        lines.append(f"  {lb:<24} {v:<14} {n}")
    lines+=["","KEY RISKS","-"*40]
    for rv in data.get("risks",[]):
        lines+=[f"  Risk: {rv.get('title','')}",f"        {rv.get('detail','')}",f"  Fix:  {rv.get('fix','')}",""]
    lines+=["RECOMMENDATIONS","-"*40]
    for i,rec in enumerate(data.get("recommendations",[]),1):
        lines+=[f"  {i}. {rec.get('action','')}",f"     {rec.get('rationale','')}",""]
    return "\n".join(lines)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG + STYLES
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="DiligenceAI", page_icon="", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
:root{--bg:#080C14;--card:#0F1620;--border:rgba(255,255,255,0.07);--bhi:rgba(79,142,247,0.35);
      --txt:#F0F4FF;--sub:#8B9BC8;--mut:#4A5578;--blue:#4F8EF7;--teal:#00D4AA;
      --purple:#9B6DFF;--amber:#F5A623;--red:#FF5C6A;}
html,body,[class*="css"]{font-family:'Inter',-apple-system,sans-serif;background:var(--bg)!important;color:var(--txt);}
.block-container{
  padding-top:1.5rem!important;
  padding-bottom:4rem!important;
  padding-left:3rem!important;
  padding-right:3rem!important;
  max-width:1280px!important;
}
#MainMenu,footer,header{visibility:hidden;}
[data-testid="stSidebarNav"],section[data-testid="stSidebar"]{display:none;}
::-webkit-scrollbar{width:6px;}::-webkit-scrollbar-track{background:var(--bg);}
::-webkit-scrollbar-thumb{background:#2A3450;border-radius:3px;}
[data-testid="stVerticalBlock"] > [data-testid="stVerticalBlock"] {gap:0.75rem;}
div[data-testid="stButton"]>button{background:rgba(255,255,255,0.04);color:var(--sub)!important;
  border:1px solid var(--border);border-radius:8px;padding:0.4rem 0.6rem;font-size:0.82rem;
  font-weight:500;width:100%;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;transition:all 0.2s;
  letter-spacing:0.2px;font-family:'Inter',sans-serif;}
div[data-testid="stButton"]>button:hover{background:rgba(79,142,247,0.15)!important;
  color:var(--blue)!important;border-color:rgba(79,142,247,0.4)!important;transform:translateY(-1px);}
div[data-testid="stButton"][data-key="analyse_btn"]>button,
div[data-testid="stButton"][data-key="screen_btn"]>button{
  background:linear-gradient(135deg,#4F8EF7,#00D4AA)!important;color:#080C14!important;
  border:none!important;font-size:0.97rem!important;font-weight:700!important;
  padding:0.75rem 2rem!important;border-radius:10px!important;
  box-shadow:0 4px 24px rgba(79,142,247,0.35)!important;}
div[data-testid="stButton"][data-key="analyse_btn"]>button:hover,
div[data-testid="stButton"][data-key="screen_btn"]>button:hover{
  transform:translateY(-2px)!important;box-shadow:0 8px 32px rgba(79,142,247,0.5)!important;}
[data-testid="metric-container"]{background:linear-gradient(135deg,#0F1620,#111827);
  border:1px solid var(--border);border-radius:12px;padding:1rem 1.1rem;transition:border-color 0.2s;}
[data-testid="metric-container"]:hover{border-color:var(--bhi);}
[data-testid="metric-container"] label{color:var(--mut)!important;font-size:0.68rem!important;
  font-weight:700!important;letter-spacing:1.2px;text-transform:uppercase;}
[data-testid="metric-container"] [data-testid="stMetricValue"]{color:var(--txt)!important;
  font-size:1.3rem!important;font-weight:700!important;}
[data-testid="metric-container"] [data-testid="stMetricDelta"]{font-size:0.72rem!important;}
[data-testid="stDownloadButton"]>button{background:rgba(255,255,255,0.03)!important;
  color:var(--sub)!important;border:1px solid var(--border)!important;border-radius:10px;
  font-size:0.85rem;font-weight:500;width:100%;white-space:nowrap;padding:0.6rem 1rem!important;
  transition:all 0.2s;font-family:'Inter',sans-serif;}
[data-testid="stDownloadButton"]>button:hover{background:rgba(79,142,247,0.1)!important;
  border-color:rgba(79,142,247,0.4)!important;color:var(--blue)!important;transform:translateY(-1px);}
[data-testid="stExpander"]{background:var(--card)!important;border:1px solid var(--border)!important;
  border-radius:10px;margin-bottom:0.5rem;}
[data-testid="stExpander"]:hover{border-color:var(--bhi)!important;}
hr{border-color:rgba(255,255,255,0.06)!important;}
textarea,input[type="text"],input[type="password"],input[type="email"]{
  background:rgba(255,255,255,0.03)!important;border:1px solid var(--border)!important;
  color:var(--txt)!important;border-radius:10px!important;font-family:'Inter',sans-serif!important;}
textarea:focus,input:focus{border-color:rgba(79,142,247,0.5)!important;
  box-shadow:0 0 0 3px rgba(79,142,247,0.1)!important;}
[data-testid="stFileUploader"]{background:rgba(255,255,255,0.02)!important;
  border:1px dashed rgba(79,142,247,0.3)!important;border-radius:12px!important;padding:0.5rem!important;}
[data-testid="stRadio"] label,[data-testid="stCheckbox"] label{color:var(--sub)!important;}
.dai-div{height:1px;background:linear-gradient(90deg,transparent,rgba(79,142,247,0.2),transparent);
  margin:2rem 0;border:none;}
[data-testid="stHorizontalBlock"]{gap:1rem!important;}
[data-testid="stAlert"]{margin-top:0.5rem!important;margin-bottom:0.5rem!important;}
[data-testid="stCaptionContainer"]{margin-top:0.2rem!important;margin-bottom:0.4rem!important;}
</style>""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────
DEFAULTS = {
    "page": "dashboard", "logged_in": False, "user_email": "", "user_id": None,
    "analysis_data": None, "loaded_analysis": None,
    "compare_ids": [], "share_id_display": None, "auth_tab": "login",
    "active_deal_id": None,
    "screening_result": None,
    "dd_result": None,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─────────────────────────────────────────────────────────────────────────────
# NAV BAR
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<div style="background:rgba(8,12,20,0.97);border-bottom:1px solid rgba(255,255,255,0.06);
            padding:0 2.5rem;display:flex;align-items:center;justify-content:center;
            height:56px;position:sticky;top:0;z-index:999;backdrop-filter:blur(20px);">
  <div style="display:flex;align-items:center;gap:0.5rem;">
    <div style="width:28px;height:28px;background:linear-gradient(135deg,#4F8EF7,#00D4AA);
                border-radius:7px;display:flex;align-items:center;justify-content:center;">
      <span style="color:#080C14;font-weight:900;font-size:0.85rem;">D</span></div>
    <span style="font-size:1.05rem;font-weight:800;color:#F0F4FF;letter-spacing:-0.4px;">DiligenceAI</span>
  </div>
</div>""", unsafe_allow_html=True)

st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)

_s, nb1, nb2, nb3, nb4, nb5, nb6, _e = st.columns([0.8, 1, 1.2, 1.2, 1, 1.5, 1, 0.8])
with nb1:
    if st.button("Dashboard", key="nb_dashboard", use_container_width=True):
        st.session_state.page = "dashboard" if st.session_state.logged_in else "account"
        st.rerun()
with nb2:
    if st.button("Deal Screening", key="nb_screening", use_container_width=True):
        st.session_state.page = "screening"
        st.rerun()
with nb3:
    if st.button("Due Diligence", key="nb_dd", use_container_width=True):
        st.session_state.page = "dd_select"
        st.rerun()
with nb4:
    if st.button("Shared", key="nb_shared", use_container_width=True):
        st.session_state.page = "shared_view"
        st.rerun()
with nb5:
    if st.button("About", key="nb_about", use_container_width=True):
        st.session_state.page = "about"
        st.rerun()
with nb6:
    if st.session_state.logged_in:
        lbl = st.session_state.user_email.split("@")[0]
        if st.button(lbl, key="nb_account", use_container_width=True):
            st.session_state.page = "account"
            st.rerun()
    else:
        if st.button("Log In / Sign Up", key="nb_account", use_container_width=True):
            st.session_state.auth_tab = "login"
            st.session_state.page = "account"
            st.rerun()

st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# SHARED UI HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def D(): st.markdown("<div class='dai-div'></div>", unsafe_allow_html=True)

def slabel(text, accent="#4F8EF7"):
    st.markdown(
        f"<div style='display:flex;align-items:center;gap:0.5rem;margin-bottom:0.8rem;'>"
        f"<div style='width:3px;height:14px;background:linear-gradient(180deg,{accent},transparent);border-radius:2px;'></div>"
        f"<span style='color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;text-transform:uppercase;'>{text}</span>"
        f"</div>", unsafe_allow_html=True)

def hcolours(label):
    return {
        "Strong":  ("rgba(0,212,170,0.08)",  "#00D4AA", "rgba(0,212,170,0.25)"),
        "Moderate":("rgba(245,166,35,0.08)", "#F5A623", "rgba(245,166,35,0.25)"),
        "Weak":    ("rgba(255,92,106,0.08)", "#FF5C6A", "rgba(255,92,106,0.25)")
    }.get(label, ("rgba(245,166,35,0.08)", "#F5A623", "rgba(245,166,35,0.25)"))

def status_colour(status):
    return {"screening": "#F5A623", "dd": "#4F8EF7", "completed": "#00D4AA"}.get(status, "#8B9BC8")

def status_label(status):
    return {"screening": "Screening", "dd": "Due Diligence", "completed": "Completed"}.get(status, status.title())

def tick(t):
    return (f"<div style='display:flex;align-items:flex-start;gap:0.7rem;padding:0.6rem 0;"
            f"border-bottom:1px solid rgba(255,255,255,0.04);'>"
            f"<span style='color:#00D4AA;font-weight:700;'>✓</span>"
            f"<span style='color:#C8D0E8;font-size:0.87rem;line-height:1.4;'>{t}</span></div>")

# ─────────────────────────────────────────────────────────────────────────────
# GROQ — SCREENING PROMPT
# ─────────────────────────────────────────────────────────────────────────────
SCREENING_PROMPT = """You are a senior investment analyst. Analyse the uploaded documents (pitch deck, CIM, financials) and return ONLY valid JSON — no markdown, no extra text.
Schema:
{"deal_name":"string","industry":"string","summary":"2-3 sentences describing what the business does",
"investment_score":0-100,
"score_label":"Strong"|"Promising"|"Borderline"|"Pass",
"recommendation":"Proceed"|"Pass",
"recommendation_rationale":"2 sentences explaining recommendation",
"key_risks":[{"title":"string","detail":"string"},{"title":"string","detail":"string"},{"title":"string","detail":"string"}],
"strengths":[{"title":"string","detail":"string"},{"title":"string","detail":"string"},{"title":"string","detail":"string"}],
"red_flags":["string","string"],
"key_metrics":{"revenue":"string or Not provided","growth_rate":"string or Not provided","gross_margin":"string or Not provided","burn_rate":"string or Not provided","runway":"string or Not provided"},
"investor_view":"2-3 sentence blunt PE-style view"}
Rules: "Not provided" for missing. Never invent figures. NZ English. Return ONLY JSON."""

# ─────────────────────────────────────────────────────────────────────────────
# GROQ — DD PROMPT
# ─────────────────────────────────────────────────────────────────────────────
DD_PROMPT = """You are a highly skilled financial analyst and forensic accountant based in New Zealand.
Cross-reference all documents together to produce a single unified analysis. Return ONLY valid JSON — no markdown, no extra text.
Schema:{"company_name":"string","period":"string","documents_detected":["list"],
"health_score":1-10,"health_label":"Strong"|"Moderate"|"Weak","health_summary":"2-3 sentences NZ English",
"kpis":{"revenue":{"value":"string","note":"string"},"net_profit":{"value":"string","note":"string"},
"gross_margin":{"value":"string","note":"string"},"net_margin":{"value":"string","note":"string"},
"ebitda":{"value":"string","note":"string"},"operating_cashflow":{"value":"string","note":"string"},
"current_ratio":{"value":"string","note":"string"},"debt_to_equity":{"value":"string","note":"string"},
"working_capital":{"value":"string","note":"string"},"total_debt":{"value":"string","note":"string"},
"revenue_growth":{"value":"string","note":"string"},"interest_coverage":{"value":"string","note":"string"}},
"profitability":{"headline":"string","points":["string","string","string"]},
"cash_health":{"headline":"string","points":["string","string","string"]},
"working_capital_analysis":{"headline":"string","points":["string","string","string"]},
"balance_sheet":{"headline":"string","points":["string","string","string"]},
"revenue_growth":{"headline":"string","points":["string","string","string"]},
"debt_leverage":{"headline":"string","points":["string","string","string"]},
"investor_view":"3-4 sentences NZ English",
"legal_flags":["string","string"],
"financial_inconsistencies":["string","string"],
"missing_information":["string","string"],
"risks":[{"title":"string","detail":"string","fix":"string"},{"title":"string","detail":"string","fix":"string"},{"title":"string","detail":"string","fix":"string"}],
"positives":[{"title":"string","detail":"string"},{"title":"string","detail":"string"},{"title":"string","detail":"string"}],
"recommendations":[{"action":"string","rationale":"string"},{"action":"string","rationale":"string"},{"action":"string","rationale":"string"}]}
Rules: "Not provided" for missing. Never invent. Format: "$12.4M","18.3%","2.1x". NZ English. Return ONLY JSON."""

def _groq_request(api_key, system_prompt, user_text):
    """Call Groq REST API directly using stdlib urllib — no SDK needed."""
    url = "https://api.groq.com/openai/v1/chat/completions"
    payload = json.dumps({
        "model": "llama-3.3-70b-versatile",
        "max_tokens": 3000,
        "temperature": 0.1,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": f"Analyse these documents:\n\n{user_text}"}
        ]
    }).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        },
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=120) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    return body["choices"][0]["message"]["content"].strip()

def call_groq(text, api_key, system_prompt):
    MAX_CHARS = 6000
    if len(text) > MAX_CHARS:
        lines = text.split("\n")
        priority = [l for l in lines if any(c.isdigit() for c in l)]
        other    = [l for l in lines if not any(c.isdigit() for c in l)]
        condensed = "\n".join(priority + other)
        if len(condensed) > MAX_CHARS:
            condensed = condensed[:MAX_CHARS]
        text = condensed + "\n\n[Document truncated to fit token limits]"
    try:
        raw = _groq_request(api_key, system_prompt, text)
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="replace")
        if e.code in (413, 429) or "rate_limit" in err_body.lower() or "too large" in err_body.lower():
            text = text[:3000] + "\n\n[Further truncated]"
            raw = _groq_request(api_key, system_prompt, text)
        else:
            raise RuntimeError(f"Groq API error {e.code}: {err_body}") from e
    raw = re.sub(r"^```(?:json)?", "", raw).strip()
    raw = re.sub(r"```$", "", raw).strip()
    try: return json.loads(raw), raw
    except Exception:
        m = re.search(r'\{.*\}', raw, re.DOTALL)
        if m:
            try: return json.loads(m.group()), raw
            except Exception: pass
    return None, raw

# ─────────────────────────────────────────────────────────────────────────────
# FILE EXTRACTORS
# ─────────────────────────────────────────────────────────────────────────────
def extract_pdf_text(file_bytes):
    try:
        import pdfplumber
        parts=[]
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages[:30]:
                t=page.extract_text()
                if t:
                    lines = [l.strip() for l in t.split("\n") if len(l.strip()) > 2]
                    if lines:
                        parts.append("\n".join(lines))
                if page.page_number <= 20:
                    for table in page.extract_tables():
                        for row in table:
                            row_text = " | ".join(str(c).strip() if c else "" for c in row)
                            if any(c.isdigit() for c in row_text):
                                parts.append(row_text)
        return "\n".join(parts)
    except Exception as e: return f"[PDF error: {e}]"

def extract_csv_text(file_bytes):
    try:
        content=file_bytes.decode("utf-8",errors="replace")
        return "\n".join(" | ".join(r) for r in csv.reader(io.StringIO(content)))
    except Exception as e: return f"[CSV error: {e}]"

def extract_excel_text(file_bytes):
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(c).strip() if c is not None else "" for c in row]
                row_text = " | ".join(row_vals)
                if any(c for c in row_vals if c):
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e: return f"[Excel error: {e}]"

def extract_file(uf):
    name=uf.name.lower(); raw=uf.read()
    if name.endswith(".pdf"): return extract_pdf_text(raw)
    if name.endswith(".csv"): return extract_csv_text(raw)
    if name.endswith((".xlsx",".xls",".xlsm")): return extract_excel_text(raw)
    return raw.decode("utf-8",errors="replace")

# ─────────────────────────────────────────────────────────────────────────────
# RENDER HELPERS  (DD analysis view)
# ─────────────────────────────────────────────────────────────────────────────
def render_banner(data):
    label=data.get("health_label","Moderate"); score=data.get("health_score",5)
    summary=data.get("health_summary",""); company=data.get("company_name","")
    period=data.get("period",""); docs=data.get("documents_detected",[])
    bg,fg,border=hcolours(label)
    bar_parts=[]
    filled=int((score/10)*20)
    for i in range(20):
        c=fg if i<filled else "rgba(255,255,255,0.08)"
        bar_parts.append(f"<span style='display:inline-block;width:8px;height:8px;border-radius:2px;margin-right:3px;background:{c};'></span>")
    bar="".join(bar_parts)
    tags="".join(f"<span style='background:rgba(79,142,247,0.1);color:#4F8EF7;border:1px solid rgba(79,142,247,0.25);border-radius:6px;padding:0.15rem 0.6rem;font-size:0.68rem;font-weight:600;margin-right:0.3rem;'>{d}</span>" for d in docs)
    st.markdown(f"""
    <div style="background:{bg};border:1px solid {border};border-radius:16px;padding:2rem 2.2rem;margin-bottom:1.5rem;position:relative;overflow:hidden;">
      <div style="position:absolute;top:0;right:0;width:200px;height:200px;background:radial-gradient(circle,{fg}08,transparent 70%);pointer-events:none;"></div>
      <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:1rem;position:relative;">
        <div>
          <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.4rem;">OVERALL FINANCIAL HEALTH</div>
          <div style="font-size:2.2rem;font-weight:800;color:{fg};line-height:1;letter-spacing:-0.5px;margin-bottom:0.3rem;">{label}</div>
          <div style="color:#8B9BC8;font-size:0.83rem;margin-bottom:0.5rem;">{company}&nbsp;·&nbsp;{period}</div>
          <div>{tags}</div>
        </div>
        <div style="text-align:right;">
          <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.3rem;">HEALTH SCORE</div>
          <div style="font-size:3rem;font-weight:900;color:{fg};line-height:1;">{score}<span style="font-size:1.2rem;color:#4A5578;font-weight:400;">/10</span></div>
          <div style="margin-top:0.5rem;">{bar}</div>
        </div>
      </div>
      <div style="margin-top:1.2rem;padding-top:1.2rem;border-top:1px solid {border};color:#C8D0E8;font-size:0.93rem;line-height:1.7;">{summary}</div>
    </div>""", unsafe_allow_html=True)

def render_kpis(kpis):
    slabel("KEY FINANCIAL METRICS")
    order=[("revenue","Revenue"),("net_profit","Net Profit"),("gross_margin","Gross Margin"),
           ("net_margin","Net Margin"),("ebitda","EBITDA"),("operating_cashflow","Operating Cash Flow"),
           ("current_ratio","Current Ratio"),("debt_to_equity","Debt / Equity"),
           ("working_capital","Working Capital"),("total_debt","Total Debt"),
           ("revenue_growth","Revenue Growth"),("interest_coverage","Interest Coverage")]
    for rs in range(0,len(order),6):
        chunk=order[rs:rs+6]; cols=st.columns(len(chunk))
        for col,(key,label) in zip(cols,chunk):
            item=kpis.get(key,{}); value=item.get("value","N/A"); note=item.get("note","")
            dc="inverse" if any(w in note.lower() for w in ["pressure","decline","weak","low","risk"]) else "normal"
            with col: st.metric(label=label,value=value,delta=note if note else None,delta_color=dc)
        st.markdown("<div style='margin-bottom:0.5rem'></div>",unsafe_allow_html=True)

def render_card(title, section, accent="#4F8EF7"):
    headline=section.get("headline",""); points=section.get("points",[])
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                border-radius:14px;padding:1.4rem 1.6rem;margin-bottom:0.8rem;border-top:2px solid {accent}22;">
      <div style="color:#F0F4FF;font-size:0.82rem;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:0.6rem;">{title}</div>
      <div style="color:#8B9BC8;font-size:0.85rem;line-height:1.55;border-left:2px solid {accent};padding-left:0.8rem;font-style:italic;">{headline}</div>
    </div>""", unsafe_allow_html=True)
    for pt in points:
        st.markdown(f"<div style='color:#C8D0E8;font-size:0.84rem;padding:0.3rem 0 0.3rem 1.1rem;border-left:2px solid rgba(255,255,255,0.06);margin-bottom:0.3rem;'>{pt}</div>", unsafe_allow_html=True)
    st.markdown("")

def render_ai_flags(data, source="dd"):
    """Render legal flags, financial inconsistencies, missing info for DD."""
    legal  = data.get("legal_flags", [])
    fincon = data.get("financial_inconsistencies", [])
    missing= data.get("missing_information", [])
    if not any([legal, fincon, missing]):
        return
    slabel("AI FLAGS & ALERTS", "#FF5C6A")
    fl1, fl2, fl3 = st.columns(3, gap="large")
    with fl1:
        st.markdown("<div style='color:#FF5C6A;font-size:0.78rem;font-weight:700;letter-spacing:1px;margin-bottom:0.6rem;'>LEGAL RISKS</div>", unsafe_allow_html=True)
        for f in legal:
            st.markdown(f"<div style='background:rgba(255,92,106,0.06);border:1px solid rgba(255,92,106,0.2);border-radius:8px;padding:0.7rem 0.9rem;margin-bottom:0.5rem;color:#C8D0E8;font-size:0.83rem;line-height:1.5;'>{f}</div>", unsafe_allow_html=True)
        if not legal:
            st.markdown("<div style='color:#4A5578;font-size:0.83rem;'>None identified.</div>", unsafe_allow_html=True)
    with fl2:
        st.markdown("<div style='color:#F5A623;font-size:0.78rem;font-weight:700;letter-spacing:1px;margin-bottom:0.6rem;'>FINANCIAL INCONSISTENCIES</div>", unsafe_allow_html=True)
        for f in fincon:
            st.markdown(f"<div style='background:rgba(245,166,35,0.06);border:1px solid rgba(245,166,35,0.2);border-radius:8px;padding:0.7rem 0.9rem;margin-bottom:0.5rem;color:#C8D0E8;font-size:0.83rem;line-height:1.5;'>{f}</div>", unsafe_allow_html=True)
        if not fincon:
            st.markdown("<div style='color:#4A5578;font-size:0.83rem;'>None identified.</div>", unsafe_allow_html=True)
    with fl3:
        st.markdown("<div style='color:#9B6DFF;font-size:0.78rem;font-weight:700;letter-spacing:1px;margin-bottom:0.6rem;'>MISSING INFORMATION</div>", unsafe_allow_html=True)
        for f in missing:
            st.markdown(f"<div style='background:rgba(155,109,255,0.06);border:1px solid rgba(155,109,255,0.2);border-radius:8px;padding:0.7rem 0.9rem;margin-bottom:0.5rem;color:#C8D0E8;font-size:0.83rem;line-height:1.5;'>{f}</div>", unsafe_allow_html=True)
        if not missing:
            st.markdown("<div style='color:#4A5578;font-size:0.83rem;'>None identified.</div>", unsafe_allow_html=True)

def render_full_analysis(data, kp="main", allow_save=True, deal_id=""):
    render_banner(data)
    render_kpis(data.get("kpis",{}))
    D()
    # AI Flags (DD-specific)
    if data.get("legal_flags") or data.get("financial_inconsistencies") or data.get("missing_information"):
        render_ai_flags(data)
        D()
    slabel("PERFORMANCE SUMMARY")
    c1,c2,c3=st.columns(3,gap="large")
    with c1: render_card("Profitability",   data.get("profitability",{}),   "#4F8EF7")
    with c2: render_card("Cash Health",     data.get("cash_health",{}),     "#00D4AA")
    with c3: render_card("Revenue Growth",  data.get("revenue_growth",{}),  "#9B6DFF")
    c4,c5,c6=st.columns(3,gap="large")
    with c4: render_card("Working Capital", data.get("working_capital_analysis",{}), "#F5A623")
    with c5: render_card("Balance Sheet",   data.get("balance_sheet",{}),   "#FF5C6A")
    with c6: render_card("Debt & Leverage", data.get("debt_leverage",{}),   "#00D4AA")
    D()
    slabel("INVESTOR VIEW")
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,rgba(79,142,247,0.06),rgba(155,109,255,0.06));
                border:1px solid rgba(79,142,247,0.2);border-radius:14px;
                padding:1.6rem 2rem;color:#C8D0E8;font-size:0.97rem;line-height:1.75;
                border-left:3px solid #4F8EF7;">{data.get("investor_view","")}</div>""",
    unsafe_allow_html=True)
    D()
    rc1,pc1=st.columns(2,gap="large")
    with rc1:
        slabel("KEY RISKS & CONCERNS","#FF5C6A")
        for risk in data.get("risks",[]):
            with st.expander(risk.get("title","Risk")):
                st.markdown(f"**Issue:** {risk.get('detail','')}")
                st.markdown(f"**Suggested Action:** {risk.get('fix','')}")
    with pc1:
        slabel("POSITIVE SIGNALS","#00D4AA")
        for pos in data.get("positives",[]):
            st.markdown(f"""
            <div style="background:rgba(0,212,170,0.06);border:1px solid rgba(0,212,170,0.2);
                        border-radius:10px;padding:0.9rem 1.1rem;margin-bottom:0.6rem;">
              <div style="color:#00D4AA;font-weight:600;font-size:0.87rem;margin-bottom:0.3rem;">{pos.get('title','')}</div>
              <div style="color:#8B9BC8;font-size:0.83rem;">{pos.get('detail','')}</div>
            </div>""", unsafe_allow_html=True)
    D()
    slabel("RECOMMENDATIONS")
    for i,rec in enumerate(data.get("recommendations",[]),1):
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                    border-left:3px solid #4F8EF7;border-radius:0 12px 12px 0;padding:1rem 1.4rem;margin-bottom:0.7rem;">
          <div style="color:#4F8EF7;font-weight:700;font-size:0.9rem;margin-bottom:0.3rem;">{i}.&nbsp;{rec.get('action','')}</div>
          <div style="color:#8B9BC8;font-size:0.84rem;line-height:1.55;">{rec.get('rationale','')}</div>
        </div>""", unsafe_allow_html=True)

    if not allow_save:
        return

    D()
    slabel("SAVE & SHARE")
    sa1,sa2,sa3,sa4=st.columns(4)
    with sa1:
        if st.session_state.logged_in and st.session_state.user_id:
            if st.button("Save Analysis", key=f"save_{kp}", use_container_width=True):
                db_save_analysis(st.session_state.user_id, data, deal_id=deal_id)
                st.success("Saved to your dashboard.")
        else:
            if st.button("Log In to Save", key=f"save_l_{kp}", use_container_width=True):
                st.session_state.page="account"; st.rerun()
    with sa2:
        if st.session_state.logged_in and st.session_state.user_id:
            if st.button("Create Share Link", key=f"share_{kp}", use_container_width=True):
                sid=db_create_share(st.session_state.user_id, data)
                st.session_state.share_id_display=sid; st.rerun()
        else:
            if st.button("Log In to Share", key=f"share_l_{kp}", use_container_width=True):
                st.session_state.page="account"; st.rerun()
    with sa3:
        if st.button("My Dashboard", key=f"dash_{kp}", use_container_width=True):
            st.session_state.page="dashboard" if st.session_state.logged_in else "account"; st.rerun()
    with sa4:
        if st.button("View Shared", key=f"vsh_{kp}", use_container_width=True):
            st.session_state.page="shared_view"; st.rerun()

    if st.session_state.share_id_display:
        sid=st.session_state.share_id_display
        st.markdown(f"""
        <div style="background:rgba(0,212,170,0.08);border:1px solid rgba(0,212,170,0.25);
                    border-radius:10px;padding:1rem 1.4rem;margin-top:0.5rem;">
          <div style="color:#00D4AA;font-weight:600;font-size:0.88rem;margin-bottom:0.3rem;">Share link created</div>
          <div style="color:#C8D0E8;font-size:0.9rem;">Share ID: <b style="font-size:1.1rem;letter-spacing:3px;">{sid}</b></div>
          <div style="color:#8B9BC8;font-size:0.8rem;margin-top:0.3rem;">Go to "Shared" in the nav bar and enter this ID.</div>
        </div>""", unsafe_allow_html=True)

    D()
    slabel("DOWNLOAD REPORT")
    slug=re.sub(r"[^a-z0-9]","_",data.get("company_name","report").lower())
    d1,d2,d3,d4=st.columns(4)
    with d1:
        st.download_button("Download (.txt)",build_txt(data),f"{slug}.txt","text/plain",use_container_width=True)
    with d2:
        if PDF_OK:
            st.download_button("Download (.pdf)",build_pdf(data),f"{slug}.pdf","application/pdf",use_container_width=True)
        else:
            st.caption("PDF library not installed")
    with d3:
        if EXCEL_OK:
            st.download_button("Download (.xlsx)",build_excel(data),f"{slug}.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
        else:
            st.caption("Excel library not installed")
    with d4:
        if DOCX_OK:
            st.download_button("Download (.docx)",build_docx(data),f"{slug}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
        else:
            st.caption("Word library not installed")

# ─────────────────────────────────────────────────────────────────────────────
# RENDER SCREENING RESULT CARD
# ─────────────────────────────────────────────────────────────────────────────
def render_screening_result(sc_data, deal=None, api_key=""):
    score       = sc_data.get("investment_score", 0)
    label       = sc_data.get("score_label", "Borderline")
    rec         = sc_data.get("recommendation", "Pass")
    rationale   = sc_data.get("recommendation_rationale", "")
    summary     = sc_data.get("summary", "")
    deal_name   = sc_data.get("deal_name", "")
    industry    = sc_data.get("industry", "")
    iv          = sc_data.get("investor_view", "")
    km          = sc_data.get("key_metrics", {})

    score_col = {"Strong":"#00D4AA","Promising":"#4F8EF7","Borderline":"#F5A623","Pass":"#FF5C6A"}.get(label,"#F5A623")
    rec_col   = "#00D4AA" if rec == "Proceed" else "#FF5C6A"
    rec_bg    = "rgba(0,212,170,0.07)" if rec == "Proceed" else "rgba(255,92,106,0.07)"
    rec_bor   = "rgba(0,212,170,0.25)" if rec == "Proceed" else "rgba(255,92,106,0.25)"

    st.markdown(f"""
    <div style="background:rgba(15,22,32,0.9);border:1px solid rgba(255,255,255,0.08);border-radius:16px;
                padding:2rem 2.4rem;margin-bottom:1.5rem;">
      <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:1rem;">
        <div>
          <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;">DEAL SCREENING RESULT</div>
          <div style="font-size:1.6rem;font-weight:800;color:#F0F4FF;margin:0.3rem 0;">{deal_name}</div>
          <div style="color:#8B9BC8;font-size:0.85rem;">{industry}</div>
        </div>
        <div style="text-align:right;">
          <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.3rem;">INVESTMENT SCORE</div>
          <div style="font-size:3rem;font-weight:900;color:{score_col};line-height:1;">{score}<span style="font-size:1rem;color:#4A5578;">/100</span></div>
          <div style="color:{score_col};font-size:0.85rem;font-weight:600;">{label}</div>
        </div>
      </div>
      <div style="margin-top:1.2rem;padding-top:1.2rem;border-top:1px solid rgba(255,255,255,0.06);
                  color:#C8D0E8;font-size:0.93rem;line-height:1.7;">{summary}</div>
    </div>""", unsafe_allow_html=True)

    # Recommendation box
    st.markdown(f"""
    <div style="background:{rec_bg};border:1px solid {rec_bor};border-radius:14px;
                padding:1.4rem 1.8rem;margin-bottom:1.5rem;display:flex;align-items:center;gap:2rem;flex-wrap:wrap;">
      <div>
        <div style="color:{rec_col};font-size:0.72rem;font-weight:700;letter-spacing:2px;margin-bottom:0.3rem;">AI RECOMMENDATION</div>
        <div style="font-size:2.2rem;font-weight:900;color:{rec_col};">{rec}</div>
      </div>
      <div style="color:#C8D0E8;font-size:0.9rem;line-height:1.6;flex:1;">{rationale}</div>
    </div>""", unsafe_allow_html=True)

    # Key metrics
    if any(v and v != "Not provided" for v in km.values()):
        slabel("KEY METRICS FROM DOCUMENTS")
        m_cols = st.columns(5)
        for col,(k,lbl) in zip(m_cols,[("revenue","Revenue"),("growth_rate","Growth Rate"),
                                       ("gross_margin","Gross Margin"),("burn_rate","Burn Rate"),("runway","Runway")]):
            with col:
                st.metric(label=lbl, value=km.get(k,"Not provided"))

    D()
    r1, r2 = st.columns(2, gap="large")
    with r1:
        slabel("KEY RISKS", "#FF5C6A")
        for risk in sc_data.get("key_risks",[]):
            with st.expander(risk.get("title","Risk")):
                st.markdown(f"{risk.get('detail','')}")
    with r2:
        slabel("STRENGTHS", "#00D4AA")
        for s in sc_data.get("strengths",[]):
            st.markdown(f"""
            <div style="background:rgba(0,212,170,0.06);border:1px solid rgba(0,212,170,0.2);
                        border-radius:10px;padding:0.9rem 1.1rem;margin-bottom:0.6rem;">
              <div style="color:#00D4AA;font-weight:600;font-size:0.87rem;margin-bottom:0.2rem;">{s.get('title','')}</div>
              <div style="color:#8B9BC8;font-size:0.82rem;">{s.get('detail','')}</div>
            </div>""", unsafe_allow_html=True)

    D()
    slabel("INVESTOR VIEW")
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,rgba(79,142,247,0.06),rgba(155,109,255,0.06));
                border:1px solid rgba(79,142,247,0.2);border-radius:14px;
                padding:1.4rem 1.8rem;color:#C8D0E8;font-size:0.93rem;line-height:1.75;
                border-left:3px solid #4F8EF7;">{iv}</div>""", unsafe_allow_html=True)

    # Red flags
    red_flags = sc_data.get("red_flags", [])
    if red_flags:
        D()
        slabel("RED FLAGS", "#FF5C6A")
        for f in red_flags:
            st.markdown(f"""
            <div style="background:rgba(255,92,106,0.06);border:1px solid rgba(255,92,106,0.2);
                        border-radius:8px;padding:0.7rem 1rem;margin-bottom:0.5rem;
                        color:#C8D0E8;font-size:0.85rem;">⚠ {f}</div>""", unsafe_allow_html=True)

    D()
    # Convert to DD button
    if st.session_state.logged_in and deal:
        status = deal.get("status","screening")
        if status == "screening":
            conv_col, _ = st.columns([1.5,2])
            with conv_col:
                if st.button("⟶  Convert to Due Diligence Project", key="convert_to_dd_btn", use_container_width=True):
                    db_promote_to_dd(deal["id"])
                    st.session_state.active_deal_id = deal["id"]
                    st.session_state.page = "dd"
                    st.rerun()
        else:
            st.markdown(f"""
            <div style="background:rgba(79,142,247,0.07);border:1px solid rgba(79,142,247,0.2);
                        border-radius:10px;padding:0.8rem 1.2rem;display:inline-flex;align-items:center;gap:0.7rem;">
              <span style="color:#4F8EF7;font-size:0.88rem;">✓ Already moved to Due Diligence</span>
            </div>""", unsafe_allow_html=True)
            if st.button("Open Due Diligence →", key="go_to_dd_btn"):
                st.session_state.active_deal_id = deal["id"]
                st.session_state.page = "dd"
                st.rerun()
    elif not st.session_state.logged_in:
        st.info("Log in to save this screening and convert it to a Due Diligence project.")

# ─────────────────────────────────────────────────────────────────────────────
# RISK TRACKER (shared between Screening results and DD)
# ─────────────────────────────────────────────────────────────────────────────
def render_risk_tracker(deal, dd_data=None):
    """Show risks from screening + any DD-identified risks, with note inputs per risk."""
    slabel("RISK TRACKER", "#F5A623")
    sc = deal.get("screening_data", {})
    screen_risks = [r.get("title","") for r in sc.get("key_risks",[])]
    dd_risks     = [r.get("title","") for r in (dd_data or {}).get("risks",[])] if dd_data else []
    # Merge, deduplicate
    all_risks = list(dict.fromkeys(screen_risks + dd_risks))

    if not all_risks:
        st.markdown("<div style='color:#4A5578;font-size:0.88rem;'>No risks identified yet. Run screening or a DD analysis first.</div>", unsafe_allow_html=True)
        return

    notes = deal.get("risk_notes", {})
    updated = False
    for risk_title in all_risks:
        source_tag = ""
        if risk_title in screen_risks and risk_title in dd_risks:
            source_tag = "<span style='background:rgba(155,109,255,0.15);color:#9B6DFF;font-size:0.68rem;padding:0.15rem 0.5rem;border-radius:4px;margin-left:0.5rem;'>Screening + DD</span>"
        elif risk_title in screen_risks:
            source_tag = "<span style='background:rgba(245,166,35,0.15);color:#F5A623;font-size:0.68rem;padding:0.15rem 0.5rem;border-radius:4px;margin-left:0.5rem;'>Screening</span>"
        else:
            source_tag = "<span style='background:rgba(79,142,247,0.15);color:#4F8EF7;font-size:0.68rem;padding:0.15rem 0.5rem;border-radius:4px;margin-left:0.5rem;'>DD</span>"

        st.markdown(f"""
        <div style="background:rgba(15,22,32,0.9);border:1px solid rgba(255,255,255,0.07);
                    border-radius:10px;padding:0.9rem 1.2rem;margin-bottom:0.5rem;">
          <div style="color:#F0F4FF;font-size:0.88rem;font-weight:600;margin-bottom:0.4rem;">
            {risk_title}{source_tag}
          </div>
        </div>""", unsafe_allow_html=True)
        note_key = f"note_{deal['id']}_{risk_title[:30].replace(' ','_')}"
        existing_note = notes.get(risk_title, "")
        new_note = st.text_area("Analyst note", value=existing_note, key=note_key,
                                placeholder="Add your notes, context, or follow-up actions...",
                                label_visibility="collapsed", height=68)
        if new_note != existing_note:
            notes[risk_title] = new_note
            updated = True

    if updated and st.session_state.logged_in:
        db_save_risk_notes(deal["id"], notes)
        st.success("Notes saved.")

# ═════════════════════════════════════════════════════════════════════════════
# ACCOUNT PAGE
# ═════════════════════════════════════════════════════════════════════════════
if st.session_state.page in ("login","signup","account"):
    if st.session_state.page == "signup":
        st.session_state.auth_tab = "signup"
    elif st.session_state.page == "login":
        st.session_state.auth_tab = "login"
    st.session_state.page = "account"

    _,ac,_=st.columns([1,1.5,1])
    with ac:
        st.markdown("""
        <div style="text-align:center;padding:1rem 0 1.5rem;">
          <div style="display:inline-flex;align-items:center;justify-content:center;width:52px;height:52px;
                      background:linear-gradient(135deg,#4F8EF7,#00D4AA);border-radius:14px;margin-bottom:1rem;">
            <span style="color:#080C14;font-weight:900;font-size:1.4rem;">D</span></div>
          <h2 style="color:#F0F4FF;font-weight:700;margin:0 0 0.3rem;">DiligenceAI</h2>
          <p style="color:#8B9BC8;font-size:0.9rem;margin:0;">Log in or create a free account to get started.</p>
        </div>""", unsafe_allow_html=True)

        t1,t2=st.columns(2)
        with t1:
            if st.button("Log In", key="tab_login", use_container_width=True):
                st.session_state.auth_tab="login"; st.rerun()
        with t2:
            if st.button("Sign Up", key="tab_signup", use_container_width=True):
                st.session_state.auth_tab="signup"; st.rerun()

        st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)

        if st.session_state.auth_tab == "login":
            email=st.text_input("Email address",placeholder="you@example.com",key="li_email")
            password=st.text_input("Password",type="password",placeholder="••••••••",key="li_pw")
            st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
            if st.button("Log In",key="login_submit",use_container_width=True):
                if email and password:
                    res=db_login_user(email,password)
                    if res["ok"]:
                        st.session_state.logged_in=True; st.session_state.user_email=res["email"]
                        st.session_state.user_id=res["user_id"]
                        st.session_state.page="dashboard"; st.rerun()
                    else: st.error(res["error"])
                else: st.error("Please enter your email and password.")
            st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
            st.markdown("<div style='text-align:center;color:#4A5578;font-size:0.83rem;'>Don't have an account? Switch to Sign Up above.</div>", unsafe_allow_html=True)
        else:
            email=st.text_input("Email address",placeholder="you@example.com",key="su_email")
            pw=st.text_input("Password",type="password",placeholder="Choose a password",key="su_pw")
            confirm=st.text_input("Confirm password",type="password",placeholder="Repeat password",key="su_cp")
            st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
            if st.button("Create Account",key="signup_submit",use_container_width=True):
                if not email or not pw: st.error("Please fill in all fields.")
                elif pw!=confirm: st.error("Passwords do not match.")
                else:
                    res=db_create_user(email,pw)
                    if res["ok"]:
                        st.session_state.logged_in=True; st.session_state.user_email=res["email"]
                        st.session_state.user_id=res["user_id"]
                        st.session_state.page="dashboard"; st.rerun()
                    else: st.error(res["error"])
            st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
            st.markdown("<div style='text-align:center;color:#4A5578;font-size:0.83rem;'>Already have an account? Switch to Log In above.</div>", unsafe_allow_html=True)

    if st.session_state.logged_in:
        st.markdown("<div style='height:1rem'></div>", unsafe_allow_html=True)
        D()
        _,lc,_ = st.columns([1,1.5,1])
        with lc:
            if st.button("Log Out", key="logout_btn", use_container_width=True):
                for k,v in DEFAULTS.items(): st.session_state[k]=v
                st.rerun()

# ═════════════════════════════════════════════════════════════════════════════
# DASHBOARD PAGE
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "dashboard":
    st.markdown("""
    <div style="padding:2rem 0 1.5rem;">
      <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.4rem;">DEAL ANALYSIS PLATFORM</div>
      <h1 style="font-size:2rem;font-weight:800;color:#F0F4FF;margin:0 0 0.3rem;">Dashboard</h1>
      <p style="color:#8B9BC8;font-size:0.95rem;margin:0;">Screen opportunities, track deals, and run deep diligence from one place.</p>
    </div>""", unsafe_allow_html=True)

    D()
    # CTA panels
    cta1, cta2 = st.columns(2, gap="large")
    with cta1:
        st.markdown("""
        <div style="background:linear-gradient(135deg,rgba(79,142,247,0.1),rgba(0,212,170,0.06));
                    border:1px solid rgba(79,142,247,0.25);border-radius:14px;padding:1.8rem 2rem;">
          <div style="color:#4F8EF7;font-size:0.72rem;font-weight:700;letter-spacing:1.5px;margin-bottom:0.5rem;">FEATURE 1</div>
          <div style="font-size:1.2rem;font-weight:700;color:#F0F4FF;margin-bottom:0.5rem;">Deal Screening</div>
          <div style="color:#8B9BC8;font-size:0.88rem;line-height:1.6;">Upload a pitch deck or CIM. AI generates an investment score, risks, strengths, and a Proceed / Pass recommendation in seconds.</div>
        </div>""", unsafe_allow_html=True)
        st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
        if st.button("Start New Deal Screening →", key="dash_to_screen", use_container_width=True):
            st.session_state.page="screening"; st.rerun()

    with cta2:
        st.markdown("""
        <div style="background:linear-gradient(135deg,rgba(155,109,255,0.1),rgba(79,142,247,0.06));
                    border:1px solid rgba(155,109,255,0.25);border-radius:14px;padding:1.8rem 2rem;">
          <div style="color:#9B6DFF;font-size:0.72rem;font-weight:700;letter-spacing:1.5px;margin-bottom:0.5rem;">FEATURE 2</div>
          <div style="font-size:1.2rem;font-weight:700;color:#F0F4FF;margin-bottom:0.5rem;">Due Diligence</div>
          <div style="color:#8B9BC8;font-size:0.88rem;line-height:1.6;">Run a deep forensic analysis across financial statements. Flag legal risks, inconsistencies, and missing data. Track risks across the full deal lifecycle.</div>
        </div>""", unsafe_allow_html=True)
        st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
        if st.button("Open Due Diligence Projects →", key="dash_to_dd", use_container_width=True):
            st.session_state.page="dd_select"; st.rerun()

    D()

    if st.session_state.logged_in and st.session_state.user_id:
        deals = db_get_deals(st.session_state.user_id)
        slabel("RECENT DEALS")
        if not deals:
            st.markdown("""
            <div style="background:#0F1620;border:1px solid rgba(255,255,255,0.07);border-radius:12px;
                        padding:2rem;text-align:center;">
              <div style="color:#8B9BC8;font-size:0.95rem;">No deals yet. Start with Deal Screening above.</div>
            </div>""", unsafe_allow_html=True)
        else:
            for deal in deals[:8]:
                sc_data   = deal.get("screening_data", {})
                d_status  = deal.get("status","screening")
                sc_label  = sc_data.get("score_label","—")
                inv_score = sc_data.get("investment_score","—")
                sc_col    = status_colour(d_status)

                card_col, act_col = st.columns([5, 1.4])
                with card_col:
                    st.markdown(f"""
                    <div style="background:#0F1620;border:1px solid rgba(255,255,255,0.07);border-radius:12px;
                                padding:1rem 1.4rem;margin-bottom:0.4rem;">
                      <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:0.5rem;">
                        <div>
                          <div style="color:#F0F4FF;font-size:0.97rem;font-weight:700;">{deal.get('name','Unnamed')}</div>
                          <div style="color:#8B9BC8;font-size:0.8rem;margin-top:0.15rem;">{deal.get('industry','') or 'No industry set'} &nbsp;·&nbsp; Updated {deal.get('updated_at','')[:10]}</div>
                        </div>
                        <div style="display:flex;gap:1.5rem;align-items:center;">
                          <div style="text-align:center;">
                            <div style="color:#8B9BC8;font-size:0.65rem;letter-spacing:1px;">STATUS</div>
                            <div style="color:{sc_col};font-weight:700;font-size:0.88rem;">{status_label(d_status)}</div>
                          </div>
                          <div style="text-align:center;">
                            <div style="color:#8B9BC8;font-size:0.65rem;letter-spacing:1px;">SCORE</div>
                            <div style="color:#4F8EF7;font-weight:700;font-size:0.88rem;">{inv_score}/100</div>
                          </div>
                        </div>
                      </div>
                    </div>""", unsafe_allow_html=True)
                with act_col:
                    st.markdown("<div style='height:0.2rem'></div>", unsafe_allow_html=True)
                    if d_status == "screening":
                        if st.button("Open Screening", key=f"open_sc_{deal['id']}", use_container_width=True):
                            st.session_state.active_deal_id = deal["id"]
                            st.session_state.screening_result = sc_data if sc_data else None
                            st.session_state.page = "screening"
                            st.rerun()
                    else:
                        if st.button("Open DD", key=f"open_dd_{deal['id']}", use_container_width=True):
                            st.session_state.active_deal_id = deal["id"]
                            st.session_state.page = "dd"
                            st.rerun()
                    if st.button("Delete", key=f"del_deal_{deal['id']}", use_container_width=True):
                        db_delete_deal(deal["id"], st.session_state.user_id); st.rerun()

        D()
        # Saved DD analyses
        analyses = db_get_analyses(st.session_state.user_id)
        if analyses:
            slabel("SAVED DUE DILIGENCE ANALYSES")
            for a in analyses:
                km=a.get("key_metrics",{}); bg,fg,border=hcolours(a.get("health_label","Moderate"))
                rev=km.get("revenue","N/A"); ebitda=km.get("ebitda","N/A")
                chk_col,card_col,act_col=st.columns([0.3,5,1.5])
                with chk_col:
                    checked=a["id"] in st.session_state.compare_ids
                    if st.checkbox("",value=checked,key=f"cmp_{a['id']}"):
                        if a["id"] not in st.session_state.compare_ids and len(st.session_state.compare_ids)<3:
                            st.session_state.compare_ids.append(a["id"])
                    else:
                        if a["id"] in st.session_state.compare_ids:
                            st.session_state.compare_ids.remove(a["id"])
                with card_col:
                    st.markdown(f"""
                    <div style="background:{bg};border:1px solid {border};border-radius:12px;padding:1rem 1.4rem;margin-bottom:0.4rem;">
                      <div style="display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:0.5rem;">
                        <div>
                          <div style="color:#F0F4FF;font-size:0.97rem;font-weight:700;">{a.get("company_name","Unknown")}</div>
                          <div style="color:#8B9BC8;font-size:0.8rem;">{a.get("period","")} &nbsp;·&nbsp; {a.get("date_created","")[:10]}</div>
                        </div>
                        <div style="display:flex;gap:1.5rem;">
                          <div style="text-align:center;"><div style="color:#8B9BC8;font-size:0.65rem;letter-spacing:1px;">REVENUE</div><div style="color:{fg};font-weight:700;font-size:0.95rem;">{rev}</div></div>
                          <div style="text-align:center;"><div style="color:#8B9BC8;font-size:0.65rem;letter-spacing:1px;">EBITDA</div><div style="color:{fg};font-weight:700;font-size:0.95rem;">{ebitda}</div></div>
                          <div style="text-align:center;"><div style="color:#8B9BC8;font-size:0.65rem;letter-spacing:1px;">HEALTH</div><div style="color:{fg};font-weight:700;font-size:0.95rem;">{a.get("health_label","—")} {a.get("health_score","—")}/10</div></div>
                        </div>
                      </div>
                    </div>""", unsafe_allow_html=True)
                with act_col:
                    if st.button("View",key=f"view_{a['id']}",use_container_width=True):
                        st.session_state.loaded_analysis=a["raw_output"]
                        st.session_state.page="view_analysis"; st.rerun()
                    if st.button("Delete",key=f"del_{a['id']}",use_container_width=True):
                        db_delete_analysis(a["id"],st.session_state.user_id); st.rerun()

            D()
            n_sel=len(st.session_state.compare_ids)
            slabel("COMPARE COMPANIES","#9B6DFF")
            st.markdown(f"<div style='color:#8B9BC8;font-size:0.88rem;margin-bottom:0.8rem;'>Tick 2–3 analyses above, then click Compare. ({n_sel} selected)</div>",unsafe_allow_html=True)
            if n_sel >= 2:
                if st.button(f"Compare {n_sel} Companies", key="compare_btn"):
                    st.session_state.page="compare"; st.rerun()
    else:
        D()
        st.markdown("""
        <div style="background:linear-gradient(135deg,rgba(79,142,247,0.08),rgba(0,212,170,0.05));
                    border:1px solid rgba(79,142,247,0.2);border-radius:14px;padding:2rem;text-align:center;">
          <div style="color:#F0F4FF;font-size:1rem;font-weight:700;margin-bottom:0.4rem;">Log in to save and track deals</div>
          <div style="color:#8B9BC8;font-size:0.88rem;margin-bottom:1rem;">Deal Screening works without an account. Log in to save results, track across the lifecycle, and run Due Diligence.</div>
        </div>""", unsafe_allow_html=True)
        st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
        _,lc,_ = st.columns([1,2,1])
        with lc:
            if st.button("Log In / Sign Up", key="dash_login", use_container_width=True):
                st.session_state.page="account"; st.rerun()

    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# DEAL SCREENING PAGE
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "screening":
    # If we have an active deal with existing screening data, show that result
    active_deal = None
    if st.session_state.active_deal_id and st.session_state.logged_in:
        active_deal = db_get_deal(st.session_state.active_deal_id)
        if active_deal and active_deal.get("screening_data"):
            st.session_state.screening_result = active_deal["screening_data"]

    if st.session_state.screening_result:
        st.markdown("""
        <div style="padding:1.5rem 0 1rem;">
          <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.3rem;">DEAL SCREENING</div>
          <h1 style="font-size:1.7rem;font-weight:800;color:#F0F4FF;margin:0;">Screening Results</h1>
        </div>""", unsafe_allow_html=True)
        if st.button("← New Screening", key="new_screen"):
            st.session_state.screening_result = None
            st.session_state.active_deal_id = None
            st.rerun()
        D()
        render_screening_result(
            st.session_state.screening_result,
            deal=active_deal,
            api_key=os.getenv("GROQ_API_KEY","")
        )
        D()
        st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)
        st.stop()

    # ── Upload form ──────────────────────────────────────────────────────────
    st.markdown("""
    <div style="text-align:center;padding:2rem 1rem 1.5rem;position:relative;">
      <div style="position:absolute;top:0;left:50%;transform:translateX(-50%);width:500px;height:160px;
                  background:radial-gradient(ellipse,rgba(79,142,247,0.1),transparent 70%);pointer-events:none;"></div>
      <div style="position:relative;">
        <div style="display:inline-block;margin-bottom:1rem;">
          <span style="background:linear-gradient(135deg,rgba(79,142,247,0.15),rgba(0,212,170,0.15));color:#4F8EF7;
                       border:1px solid rgba(79,142,247,0.3);border-radius:20px;padding:0.3rem 1rem;
                       font-size:0.72rem;font-weight:700;letter-spacing:1.5px;">AI-POWERED DEAL SCREENING</span></div>
        <h1 style="font-size:2.4rem;font-weight:900;line-height:1.1;letter-spacing:-0.8px;margin:0 0 1rem;">
          <span style="color:#F0F4FF;">Screen any deal</span><br>
          <span style="background:linear-gradient(135deg,#4F8EF7,#00D4AA);-webkit-background-clip:text;
                       -webkit-text-fill-color:transparent;background-clip:text;">in under a minute.</span></h1>
        <p style="color:#8B9BC8;font-size:1rem;max-width:520px;margin:0 auto;line-height:1.7;">
          Upload a pitch deck or CIM. AI generates a structured investment assessment with score, risks, strengths, and a clear recommendation.</p>
      </div>
    </div>""", unsafe_allow_html=True)

    api_key = os.getenv("GROQ_API_KEY","")
    if not api_key:
        with st.expander("Enter Groq API Key  —  free at console.groq.com"):
            api_key = st.text_input("Key",type="password",placeholder="gsk_...",label_visibility="collapsed")
            st.caption("Get your free key at [console.groq.com](https://console.groq.com) → API Keys → Create Key")
        if not api_key: st.info("Enter your Groq API key above to get started.")

    D()
    cl, cr = st.columns(2, gap="large")
    with cl:
        st.markdown("<div style='color:#F0F4FF;font-size:0.9rem;font-weight:600;margin-bottom:0.4rem;'>Deal Name</div>",unsafe_allow_html=True)
        deal_name_input = st.text_input("Deal name",placeholder="e.g. Apex FinTech Series B",label_visibility="collapsed")
        st.markdown("<div style='color:#F0F4FF;font-size:0.9rem;font-weight:600;margin-bottom:0.4rem;margin-top:0.8rem;'>Industry (optional)</div>",unsafe_allow_html=True)
        industry_input = st.text_input("Industry",placeholder="e.g. FinTech, HealthTech, SaaS",label_visibility="collapsed")
    with cr:
        st.markdown("<div style='color:#F0F4FF;font-size:0.9rem;font-weight:600;margin-bottom:0.4rem;'>Upload Documents <span style='color:#4A5578;font-size:0.8rem;'>— Pitch deck, CIM, financials</span></div>",unsafe_allow_html=True)
        st.caption("PDF, PPT, Excel, CSV, TXT accepted")
        sc_files = st.file_uploader("sc_files",type=["pdf","csv","xlsx","xls","pptx","docx","txt"],
                                     accept_multiple_files=True, label_visibility="collapsed")
        if sc_files:
            for f in sc_files:
                kb = len(f.getvalue())/1024
                st.markdown(f"<div style='display:flex;align-items:center;gap:0.5rem;padding:0.35rem 0.6rem;background:rgba(79,142,247,0.06);border:1px solid rgba(79,142,247,0.15);border-radius:6px;margin-top:0.3rem;'><span style='color:#4F8EF7;font-size:0.75rem;'>DOC</span><span style='color:#C8D0E8;font-size:0.82rem;'>{f.name}</span><span style='color:#4A5578;font-size:0.75rem;margin-left:auto;'>{kb:.1f} KB</span></div>",unsafe_allow_html=True)

    st.markdown("<div style='height:0.5rem'></div>", unsafe_allow_html=True)
    st.markdown("<div style='color:#F0F4FF;font-size:0.9rem;font-weight:600;margin-bottom:0.4rem;'>Or Paste Document Text <span style='color:#4A5578;font-size:0.8rem;'>— any format</span></div>",unsafe_allow_html=True)
    sc_pasted = st.text_area("Paste text",height=120,placeholder="Paste pitch deck text, financials, or any deal information...",label_visibility="collapsed")

    D()
    _,sc_btn_col,_ = st.columns([1,2,1])
    with sc_btn_col:
        go_screen = st.button("Run AI Screening", key="screen_btn", use_container_width=True)

    if go_screen:
        if not api_key: st.error("Please enter your Groq API key."); st.stop()
        if not deal_name_input.strip(): st.error("Please enter a deal name."); st.stop()
        parts = []
        if sc_files:
            for uf in sc_files: uf.seek(0); parts.append(f"=== DOCUMENT: {uf.name} ===\n{extract_file(uf)}")
        if sc_pasted.strip(): parts.append(f"=== PASTED DATA ===\n{sc_pasted.strip()}")
        if not parts: st.warning("Please upload a file or paste deal information."); st.stop()

        context = f"Deal Name: {deal_name_input}\nIndustry: {industry_input}\n\n" + "\n\n".join(parts)
        with st.spinner("Screening deal with AI..."):
            try:
                sc_data, sc_raw = call_groq(context, api_key, SCREENING_PROMPT)
            except Exception as e:
                err = str(e).lower()
                if "401" in err or "invalid api key" in err:
                    st.error("Invalid Groq API key.")
                elif "429" in err or "rate_limit" in err:
                    st.error("Groq rate limit reached. Wait 60 seconds and try again.")
                else:
                    st.error(f"API error: {e}")
                st.stop()

        if not sc_data:
            st.warning("Could not parse AI output."); st.text(sc_raw); st.stop()

        # Patch deal name / industry if AI didn't fill them
        if not sc_data.get("deal_name") or sc_data["deal_name"] in ("string",""):
            sc_data["deal_name"] = deal_name_input
        if not sc_data.get("industry") or sc_data["industry"] in ("string",""):
            sc_data["industry"] = industry_input or "Not specified"

        file_names = [f.name for f in sc_files] if sc_files else []
        st.session_state.screening_result = sc_data

        # Save deal to DB if logged in
        if st.session_state.logged_in and st.session_state.user_id:
            deal_id = db_create_deal(
                st.session_state.user_id,
                deal_name_input,
                industry_input,
                file_names
            )
            db_save_screening(deal_id, sc_data, file_names)
            st.session_state.active_deal_id = deal_id

        st.rerun()

    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# DUE DILIGENCE — SELECT DEAL
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "dd_select":
    st.markdown("""
    <div style="padding:2rem 0 1.5rem;">
      <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.4rem;">DUE DILIGENCE</div>
      <h1 style="font-size:1.8rem;font-weight:800;color:#F0F4FF;margin:0 0 0.3rem;">Select a Deal</h1>
      <p style="color:#8B9BC8;font-size:0.95rem;margin:0;">Choose an existing deal to open in DD, or start fresh with a new screening.</p>
    </div>""", unsafe_allow_html=True)
    D()

    if not st.session_state.logged_in:
        st.info("Log in to access your deals and Due Diligence projects.")
        if st.button("Log In / Sign Up", key="dd_login"):
            st.session_state.page="account"; st.rerun()
    else:
        deals = db_get_deals(st.session_state.user_id)
        dd_deals = [d for d in deals if d.get("status") in ("dd","completed")]
        sc_only  = [d for d in deals if d.get("status") == "screening"]

        if dd_deals:
            slabel("ACTIVE DUE DILIGENCE DEALS")
            for deal in dd_deals:
                c, b = st.columns([5, 1.5])
                with c:
                    sc_data = deal.get("screening_data", {})
                    sc_col  = status_colour(deal.get("status","dd"))
                    st.markdown(f"""
                    <div style="background:#0F1620;border:1px solid rgba(255,255,255,0.07);border-radius:12px;
                                padding:1rem 1.4rem;margin-bottom:0.4rem;">
                      <div style="color:#F0F4FF;font-size:0.97rem;font-weight:700;">{deal.get('name','Unnamed')}</div>
                      <div style="color:#8B9BC8;font-size:0.8rem;margin-top:0.2rem;">{deal.get('industry','') or '—'} &nbsp;·&nbsp; <span style="color:{sc_col};">{status_label(deal.get('status','dd'))}</span> &nbsp;·&nbsp; Score: {sc_data.get('investment_score','—')}/100</div>
                    </div>""", unsafe_allow_html=True)
                with b:
                    st.markdown("<div style='height:0.2rem'></div>", unsafe_allow_html=True)
                    if st.button("Open DD", key=f"sel_dd_{deal['id']}", use_container_width=True):
                        st.session_state.active_deal_id = deal["id"]
                        st.session_state.page = "dd"
                        st.rerun()

        if sc_only:
            D()
            slabel("SCREENING STAGE — CONVERT TO PROCEED", "#F5A623")
            for deal in sc_only:
                c, b = st.columns([5, 1.5])
                with c:
                    sc_data = deal.get("screening_data", {})
                    st.markdown(f"""
                    <div style="background:#0F1620;border:1px solid rgba(245,166,35,0.2);border-radius:12px;
                                padding:1rem 1.4rem;margin-bottom:0.4rem;">
                      <div style="color:#F0F4FF;font-size:0.97rem;font-weight:700;">{deal.get('name','Unnamed')}</div>
                      <div style="color:#8B9BC8;font-size:0.8rem;margin-top:0.2rem;">{deal.get('industry','') or '—'} &nbsp;·&nbsp; <span style="color:#F5A623;">Screening</span> &nbsp;·&nbsp; Score: {sc_data.get('investment_score','—')}/100</div>
                    </div>""", unsafe_allow_html=True)
                with b:
                    st.markdown("<div style='height:0.2rem'></div>", unsafe_allow_html=True)
                    if st.button("Promote to DD", key=f"promote_{deal['id']}", use_container_width=True):
                        db_promote_to_dd(deal["id"])
                        st.session_state.active_deal_id = deal["id"]
                        st.session_state.page = "dd"
                        st.rerun()

        if not deals:
            st.markdown("""
            <div style="background:#0F1620;border:1px solid rgba(255,255,255,0.07);border-radius:12px;
                        padding:2rem;text-align:center;">
              <div style="color:#8B9BC8;font-size:0.95rem;">No deals yet. Run a Deal Screening first.</div>
            </div>""", unsafe_allow_html=True)

        D()
        if st.button("+ New Deal Screening", key="dd_to_screen"):
            st.session_state.page="screening"; st.rerun()

# ═════════════════════════════════════════════════════════════════════════════
# DUE DILIGENCE — MAIN WORKSPACE
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "dd":
    if not st.session_state.active_deal_id:
        st.session_state.page = "dd_select"; st.rerun()

    deal = db_get_deal(st.session_state.active_deal_id) if st.session_state.logged_in else None
    if not deal and st.session_state.logged_in:
        st.session_state.page = "dd_select"; st.rerun()

    sc_data = deal.get("screening_data", {}) if deal else {}
    dd_data = deal.get("dd_data", {}) if deal else {}

    st.markdown("<div style='margin-bottom:0.3rem'>", unsafe_allow_html=True)
    if st.button("← Back to Deal List", key="back_dd_list"):
        st.session_state.page = "dd_select"; st.rerun()

    deal_name = deal.get("name","Deal") if deal else "Due Diligence"
    st.markdown(f"""
    <div style="padding:1rem 0 0.5rem;">
      <div style="color:#8B9BC8;font-size:0.7rem;font-weight:700;letter-spacing:2px;margin-bottom:0.3rem;">DUE DILIGENCE WORKSPACE</div>
      <h1 style="font-size:1.8rem;font-weight:800;color:#F0F4FF;margin:0 0 0.3rem;">{deal_name}</h1>
      <div style="color:{status_colour(deal.get('status','dd') if deal else 'dd')};font-size:0.85rem;font-weight:600;">{status_label(deal.get('status','dd') if deal else 'dd')} &nbsp;·&nbsp; <span style="color:#8B9BC8;">{deal.get('industry','') if deal else ''}</span></div>
    </div>""", unsafe_allow_html=True)

    # If there's screening data, show import banner
    if sc_data and sc_data.get("investment_score"):
        sc_col = "#00D4AA" if sc_data.get("recommendation")=="Proceed" else "#FF5C6A"
        st.markdown(f"""
        <div style="background:rgba(79,142,247,0.07);border:1px solid rgba(79,142,247,0.2);
                    border-radius:10px;padding:0.8rem 1.2rem;margin:0.5rem 0 1rem;
                    display:flex;align-items:center;gap:1rem;flex-wrap:wrap;">
          <span style="color:#4F8EF7;font-size:0.85rem;font-weight:600;">✓ Screening data imported</span>
          <span style="color:#8B9BC8;font-size:0.83rem;">Score: <b style="color:#4F8EF7;">{sc_data.get('investment_score')}/100</b> &nbsp;·&nbsp; Recommendation: <b style="color:{sc_col};">{sc_data.get('recommendation','—')}</b></span>
        </div>""", unsafe_allow_html=True)

    D()

    api_key = os.getenv("GROQ_API_KEY","")
    if not api_key:
        with st.expander("Enter Groq API Key  —  free at console.groq.com"):
            api_key = st.text_input("Key",type="password",placeholder="gsk_...",label_visibility="collapsed",key="dd_api")
            st.caption("Get your free key at [console.groq.com](https://console.groq.com) → API Keys → Create Key")

    # DD upload
    slabel("UPLOAD FINANCIAL STATEMENTS")
    st.caption("Upload Income Statement, Balance Sheet, and Cash Flow Statement for forensic analysis. Files from screening are pre-loaded conceptually — re-upload for full DD analysis.")
    cl2, cr2 = st.columns(2, gap="large")
    with cl2:
        if deal and deal.get("file_names"):
            st.markdown("<div style='color:#8B9BC8;font-size:0.82rem;margin-bottom:0.5rem;'>Files from screening:</div>",unsafe_allow_html=True)
            for fn in deal["file_names"]:
                st.markdown(f"<div style='display:flex;align-items:center;gap:0.5rem;padding:0.3rem 0.6rem;background:rgba(79,142,247,0.05);border:1px solid rgba(79,142,247,0.12);border-radius:6px;margin-bottom:0.25rem;'><span style='color:#4F8EF7;font-size:0.75rem;'>DOC</span><span style='color:#C8D0E8;font-size:0.82rem;'>{fn}</span></div>",unsafe_allow_html=True)
    with cr2:
        st.markdown("<div style='color:#F0F4FF;font-size:0.88rem;font-weight:600;margin-bottom:0.4rem;'>Upload additional documents</div>",unsafe_allow_html=True)
        dd_files = st.file_uploader("dd_files",type=["pdf","csv","xlsx","xls","docx","txt"],
                                     accept_multiple_files=True,label_visibility="collapsed",key="dd_upload")
        if dd_files:
            for f in dd_files:
                kb=len(f.getvalue())/1024
                st.markdown(f"<div style='display:flex;align-items:center;gap:0.5rem;padding:0.35rem 0.6rem;background:rgba(79,142,247,0.06);border:1px solid rgba(79,142,247,0.15);border-radius:6px;margin-top:0.3rem;'><span style='color:#4F8EF7;font-size:0.75rem;'>DOC</span><span style='color:#C8D0E8;font-size:0.82rem;'>{f.name}</span><span style='color:#4A5578;font-size:0.75rem;margin-left:auto;'>{kb:.1f} KB</span></div>",unsafe_allow_html=True)

    dd_pasted = st.text_area("Or paste financial data",height=100,placeholder="Paste financial statement data...",label_visibility="visible",key="dd_paste")

    D()
    _,dd_btn,_ = st.columns([1,2,1])
    with dd_btn:
        go_dd = st.button("Run Deep Due Diligence Analysis", key="analyse_btn", use_container_width=True)

    if go_dd:
        if not api_key: st.error("Please enter your Groq API key."); st.stop()
        parts = []
        # Prepend screening summary as context
        if sc_data:
            parts.append(f"=== SCREENING SUMMARY (pre-imported) ===\nDeal: {sc_data.get('deal_name','')}\nSummary: {sc_data.get('summary','')}\nInvestment Score: {sc_data.get('investment_score','')}/100\nRecommendation: {sc_data.get('recommendation','')}")
        if dd_files:
            for uf in dd_files: uf.seek(0); parts.append(f"=== DOCUMENT: {uf.name} ===\n{extract_file(uf)}")
        if dd_pasted.strip(): parts.append(f"=== PASTED DATA ===\n{dd_pasted.strip()}")
        if len(parts) <= 1 and not dd_files and not dd_pasted.strip():
            st.warning("Please upload financial statements or paste data for the DD analysis."); st.stop()

        with st.spinner("Running deep due diligence analysis..."):
            try:
                dd_result, dd_raw = call_groq("\n\n".join(parts), api_key, DD_PROMPT)
            except Exception as e:
                err=str(e).lower()
                if "401" in err or "invalid api key" in err: st.error("Invalid Groq API key.")
                elif "429" in err or "rate_limit" in err: st.error("Rate limit reached. Wait 60s and retry.")
                else: st.error(f"API error: {e}")
                st.stop()

        if not dd_result:
            st.warning("Could not parse output."); st.text(dd_raw); st.stop()

        # Persist
        if deal:
            db_save_dd(deal["id"], dd_result)
        st.session_state.dd_result = dd_result
        st.rerun()

    # Show existing DD result
    if st.session_state.dd_result or dd_data:
        result_to_show = st.session_state.dd_result or dd_data
        D()
        st.markdown("""
        <div style="background:rgba(0,212,170,0.07);border:1px solid rgba(0,212,170,0.2);border-radius:10px;
                    padding:0.7rem 1.1rem;margin-bottom:1rem;display:inline-flex;align-items:center;gap:0.6rem;">
          <span style="color:#00D4AA;font-size:0.88rem;font-weight:600;">✓ Due Diligence analysis complete</span>
        </div>""", unsafe_allow_html=True)
        render_full_analysis(result_to_show, kp="dd_main", allow_save=True,
                             deal_id=deal["id"] if deal else "")

    # Risk tracker always shown in DD
    if deal:
        D()
        render_risk_tracker(deal, dd_data=result_to_show if (st.session_state.dd_result or dd_data) else None)

    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# VIEW SAVED ANALYSIS
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "view_analysis":
    data=st.session_state.loaded_analysis
    if not data: st.session_state.page="dashboard"; st.rerun()
    if st.button("← Back to Dashboard",key="back_view"): st.session_state.page="dashboard"; st.rerun()
    D()
    render_full_analysis(data, kp="view", allow_save=True)
    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# COMPARE PAGE
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "compare":
    if st.button("← Back to Dashboard",key="back_compare"): st.session_state.page="dashboard"; st.rerun()
    ids=st.session_state.compare_ids
    if len(ids)<2: st.session_state.page="dashboard"; st.rerun()
    loaded=[db_get_analysis(aid) for aid in ids if db_get_analysis(aid)]
    companies=[a.get("raw_output",{}) for a in loaded]
    st.markdown("""
    <div style="padding:2rem 0 1.5rem;">
      <h1 style="font-size:1.8rem;font-weight:800;color:#F0F4FF;margin:0 0 0.3rem;">Company Comparison</h1>
      <p style="color:#8B9BC8;margin:0;">Side-by-side metrics for selected companies.</p>
    </div>""", unsafe_allow_html=True)
    D()
    KPI_ROWS=[("revenue","Revenue"),("net_profit","Net Profit"),("gross_margin","Gross Margin"),
              ("net_margin","Net Margin"),("ebitda","EBITDA"),("operating_cashflow","Operating Cash Flow"),
              ("current_ratio","Current Ratio"),("debt_to_equity","Debt / Equity"),
              ("working_capital","Working Capital"),("total_debt","Total Debt")]
    hdr_cols=st.columns([2]+[1]*len(companies))
    hdr_cols[0].markdown("<div style='color:#8B9BC8;font-size:0.72rem;font-weight:700;letter-spacing:1px;padding:0.5rem 0;'>METRIC</div>",unsafe_allow_html=True)
    for i,comp in enumerate(companies):
        bg,fg,border=hcolours(comp.get("health_label","Moderate"))
        hdr_cols[i+1].markdown(f"""
        <div style="background:{bg};border:1px solid {border};border-radius:10px;padding:0.7rem;text-align:center;margin-bottom:0.3rem;">
          <div style="color:#F0F4FF;font-size:0.87rem;font-weight:700;">{comp.get('company_name','Unknown')}</div>
          <div style="color:{fg};font-size:0.75rem;font-weight:600;">{comp.get('health_label','—')} · {comp.get('health_score','—')}/10</div>
          <div style="color:#8B9BC8;font-size:0.7rem;">{comp.get('period','')}</div>
        </div>""", unsafe_allow_html=True)
    for key,label in KPI_ROWS:
        vals=[c.get("kpis",{}).get(key,{}).get("value","N/A") for c in companies]
        row_cols=st.columns([2]+[1]*len(companies))
        row_cols[0].markdown(f"<div style='color:#8B9BC8;font-size:0.85rem;padding:0.5rem 0;border-top:1px solid rgba(255,255,255,0.04);'>{label}</div>",unsafe_allow_html=True)
        for i,val in enumerate(vals):
            row_cols[i+1].markdown(f"<div style='color:#F0F4FF;font-size:0.9rem;font-weight:600;padding:0.5rem 0;border-top:1px solid rgba(255,255,255,0.04);text-align:center;'>{val}</div>",unsafe_allow_html=True)
    D()
    slabel("HEALTH SUMMARIES")
    sum_cols=st.columns(len(companies))
    for col,comp in zip(sum_cols,companies):
        bg,fg,border=hcolours(comp.get("health_label","Moderate"))
        with col:
            st.markdown(f"""
            <div style="background:{bg};border:1px solid {border};border-radius:12px;padding:1.2rem;">
              <div style="color:{fg};font-size:0.87rem;font-weight:700;margin-bottom:0.5rem;">{comp.get('company_name','Unknown')}</div>
              <div style="color:#C8D0E8;font-size:0.83rem;line-height:1.6;">{comp.get('health_summary','')}</div>
            </div>""", unsafe_allow_html=True)
    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# SHARED VIEW
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "shared_view":
    st.markdown("""
    <div style="padding:2rem 0 1.5rem;">
      <h1 style="font-size:1.8rem;font-weight:800;color:#F0F4FF;margin:0 0 0.3rem;">View Shared Report</h1>
      <p style="color:#8B9BC8;margin:0;">Enter a Share ID to view a shared analysis.</p>
    </div>""", unsafe_allow_html=True)
    D()
    _,mid,_=st.columns([1,2,1])
    with mid:
        sid_input=st.text_input("Share ID",placeholder="e.g. A1B2C3D4",key="sid_input")
        if st.button("Load Report",key="load_share",use_container_width=True):
            if sid_input.strip():
                shared=db_get_shared(sid_input.strip())
                if shared:
                    st.session_state.loaded_analysis=shared["raw_output"]
                    st.session_state.page="shared_display"; st.rerun()
                else: st.error("Share ID not found.")
            else: st.error("Please enter a Share ID.")

elif st.session_state.page == "shared_display":
    data=st.session_state.loaded_analysis
    if not data: st.session_state.page="shared_view"; st.rerun()
    if st.button("← Enter a different Share ID",key="back_share"): st.session_state.page="shared_view"; st.rerun()
    st.markdown("""
    <div style="background:rgba(79,142,247,0.08);border:1px solid rgba(79,142,247,0.2);border-radius:10px;
                padding:0.7rem 1.2rem;margin-bottom:1rem;">
      <span style="color:#4F8EF7;font-size:0.85rem;">Read-only shared report.</span>
    </div>""", unsafe_allow_html=True)
    D()
    render_full_analysis(data, kp="shared", allow_save=False)
    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# ABOUT PAGE
# ═════════════════════════════════════════════════════════════════════════════
elif st.session_state.page == "about":
    st.markdown("""
    <div style="text-align:center;padding:2.5rem 1rem 2rem;position:relative;">
      <div style="position:absolute;top:0;left:50%;transform:translateX(-50%);width:500px;height:160px;
                  background:radial-gradient(ellipse,rgba(155,109,255,0.1),transparent 70%);pointer-events:none;"></div>
      <div style="position:relative;">
        <div style="display:inline-block;margin-bottom:1rem;">
          <span style="background:linear-gradient(135deg,rgba(155,109,255,0.15),rgba(79,142,247,0.15));color:#9B6DFF;
                       border:1px solid rgba(155,109,255,0.3);border-radius:20px;padding:0.3rem 1rem;
                       font-size:0.72rem;font-weight:700;letter-spacing:1.5px;">DEAL ANALYSIS PLATFORM</span></div>
        <h1 style="font-size:2.5rem;font-weight:900;line-height:1.1;letter-spacing:-0.8px;margin:0 0 1rem;color:#F0F4FF;">
          Institutional-grade analysis.<br>
          <span style="background:linear-gradient(135deg,#9B6DFF,#4F8EF7);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">Available to everyone.</span></h1>
        <p style="color:#8B9BC8;font-size:1rem;max-width:600px;margin:0 auto;line-height:1.7;">
          DiligenceAI combines Deal Screening and Due Diligence into one connected workflow. Upload once, track the whole lifecycle.</p>
      </div>
    </div>""", unsafe_allow_html=True)

    D()

    # ── Deal Screening Section ───────────────────────────────────────────────
    slabel("DEAL SCREENING")
    sc1, sc2 = st.columns([1.1, 1], gap="large")
    with sc1:
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                    border-top:2px solid #4F8EF7;border-radius:14px;padding:1.8rem 2rem;">
          <div style="font-size:1.15rem;font-weight:700;color:#F0F4FF;margin-bottom:0.8rem;">What is Deal Screening?</div>
          <div style="color:#C8D0E8;font-size:0.93rem;line-height:1.75;margin-bottom:1.2rem;">
            Deal Screening uses AI to quickly evaluate investment opportunities by analysing pitch decks, financials, and other documents to identify key risks, strengths, and overall attractiveness.
          </div>
          <div style="color:#8B9BC8;font-size:0.72rem;font-weight:700;letter-spacing:1.5px;margin-bottom:0.7rem;">WHAT THE AI DOES</div>
          {''.join([f"<div style='display:flex;align-items:flex-start;gap:0.7rem;padding:0.5rem 0;border-bottom:1px solid rgba(255,255,255,0.04);'><span style='color:#00D4AA;font-weight:700;margin-top:0.05rem;'>✓</span><span style='color:#C8D0E8;font-size:0.87rem;line-height:1.4;'>{t}</span></div>" for t in [
            "Upload pitch decks or CIMs (PDF, PPT, Excel, Word)",
            "AI generates a business summary and investment score (0–100)",
            "Identifies key risks automatically",
            "Highlights strengths and competitive advantages",
            "Recommends <b style='color:#00D4AA;'>Proceed</b> or <b style='color:#FF5C6A;'>Pass</b>"
          ]])}
        </div>""", unsafe_allow_html=True)
    with sc2:
        st.markdown("""
        <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                    border-top:2px solid #00D4AA;border-radius:14px;padding:1.8rem 2rem;">
          <div style="font-size:1.15rem;font-weight:700;color:#F0F4FF;margin-bottom:0.8rem;">How It Works</div>
          <div style="display:flex;flex-direction:column;gap:0.8rem;">""", unsafe_allow_html=True)

        for i,(num,title,desc) in enumerate([
            ("01","Upload documents","Drop in a pitch deck, CIM, or financial model — PDF, PPT, Excel all supported."),
            ("02","AI extracts & analyses","The AI reads all documents, cross-references data, and identifies key patterns."),
            ("03","Structured insights output","You receive a score, summary, risks, strengths, and a clear Proceed / Pass recommendation."),
            ("04","Convert to Due Diligence","If you proceed, one click transfers all context into the full DD workspace — no re-uploading.")
        ]):
            st.markdown(f"""
            <div style="display:flex;gap:1rem;align-items:flex-start;padding:0.8rem 0;border-bottom:1px solid rgba(255,255,255,0.04);">
              <div style="background:linear-gradient(135deg,#4F8EF7,#00D4AA);-webkit-background-clip:text;
                          -webkit-text-fill-color:transparent;background-clip:text;font-size:1.5rem;font-weight:900;
                          flex-shrink:0;width:32px;">{num}</div>
              <div>
                <div style="color:#F0F4FF;font-size:0.88rem;font-weight:600;margin-bottom:0.2rem;">{title}</div>
                <div style="color:#8B9BC8;font-size:0.82rem;line-height:1.5;">{desc}</div>
              </div>
            </div>""", unsafe_allow_html=True)

        st.markdown("</div></div>", unsafe_allow_html=True)

    D()
    # ── Due Diligence Section ────────────────────────────────────────────────
    slabel("DUE DILIGENCE")
    dd1, dd2 = st.columns([1.1, 1], gap="large")
    with dd1:
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                    border-top:2px solid #9B6DFF;border-radius:14px;padding:1.8rem 2rem;">
          <div style="font-size:1.15rem;font-weight:700;color:#F0F4FF;margin-bottom:0.8rem;">Deep Forensic Analysis</div>
          <div style="color:#C8D0E8;font-size:0.93rem;line-height:1.75;margin-bottom:1.2rem;">
            Once a deal passes screening, move it to Due Diligence for a full chartered-accountant-style forensic review. All context from screening is automatically imported.
          </div>
          <div style="color:#8B9BC8;font-size:0.72rem;font-weight:700;letter-spacing:1.5px;margin-bottom:0.7rem;">AI CAPABILITIES IN DD</div>
          {''.join([f"<div style='display:flex;align-items:flex-start;gap:0.7rem;padding:0.5rem 0;border-bottom:1px solid rgba(255,255,255,0.04);'><span style='color:#9B6DFF;font-weight:700;margin-top:0.05rem;'>✓</span><span style='color:#C8D0E8;font-size:0.87rem;line-height:1.4;'>{t}</span></div>" for t in [
            "Forensic review of Income Statement, Balance Sheet, and Cash Flow",
            "Flags legal risks including change of control clauses",
            "Detects financial inconsistencies across documents",
            "Identifies missing information that should be requested",
            "12-metric KPI analysis with commentary",
            "Full risk, positives, and recommendations report"
          ]])}
        </div>""", unsafe_allow_html=True)
    with dd2:
        st.markdown("""
        <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                    border-top:2px solid #F5A623;border-radius:14px;padding:1.8rem 2rem;">
          <div style="font-size:1.15rem;font-weight:700;color:#F0F4FF;margin-bottom:0.8rem;">Risk Tracking</div>
          <div style="color:#C8D0E8;font-size:0.9rem;line-height:1.7;margin-bottom:1rem;">
            Risks identified during screening automatically appear in the DD Risk Tracker. As DD uncovers additional risks, they're added and tagged by source.
          </div>
          <div style="background:rgba(245,166,35,0.06);border:1px solid rgba(245,166,35,0.2);border-radius:10px;padding:1rem 1.2rem;margin-bottom:0.8rem;">
            <div style="color:#F5A623;font-size:0.85rem;font-weight:600;margin-bottom:0.3rem;">From Screening</div>
            <div style="color:#8B9BC8;font-size:0.82rem;">Risks carry forward with their original context</div>
          </div>
          <div style="background:rgba(79,142,247,0.06);border:1px solid rgba(79,142,247,0.2);border-radius:10px;padding:1rem 1.2rem;margin-bottom:0.8rem;">
            <div style="color:#4F8EF7;font-size:0.85rem;font-weight:600;margin-bottom:0.3rem;">From DD Analysis</div>
            <div style="color:#8B9BC8;font-size:0.82rem;">New risks surfaced during deep analysis are tagged separately</div>
          </div>
          <div style="background:rgba(0,212,170,0.06);border:1px solid rgba(0,212,170,0.2);border-radius:10px;padding:1rem 1.2rem;">
            <div style="color:#00D4AA;font-size:0.85rem;font-weight:600;margin-bottom:0.3rem;">Analyst Notes</div>
            <div style="color:#8B9BC8;font-size:0.82rem;">Add notes, actions, and context per risk — saved to the deal record</div>
          </div>
        </div>""", unsafe_allow_html=True)

    D()
    # ── Who is it for ────────────────────────────────────────────────────────
    slabel("WHO IS IT FOR")
    fc1,fc2,fc3,fc4 = st.columns(4, gap="medium")
    for col,(title,desc,colour) in zip([fc1,fc2,fc3,fc4],[
        ("Investors","Screen and track deals across the full investment lifecycle.",    "#4F8EF7"),
        ("PE / VC","Run institutional-quality forensic analysis at the speed of AI.", "#00D4AA"),
        ("Accountants","Generate structured analysis for clients in seconds.",          "#9B6DFF"),
        ("Business Owners","Understand your financials before a capital raise.",        "#F5A623")]):
        with col:
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,#0F1620,#111827);border:1px solid rgba(255,255,255,0.07);
                        border-radius:12px;padding:1.4rem;margin-bottom:0.8rem;border-top:2px solid {colour}33;">
              <div style="color:{colour};font-size:0.88rem;font-weight:700;margin-bottom:0.5rem;">{title}</div>
              <div style="color:#8B9BC8;font-size:0.83rem;line-height:1.6;">{desc}</div>
            </div>""", unsafe_allow_html=True)

    D()
    st.markdown("<p style='text-align:center;color:#4A5578;font-size:0.76rem;'>DiligenceAI &nbsp;·&nbsp; For informational purposes only — not financial advice.</p>",unsafe_allow_html=True)
